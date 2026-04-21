import io
import base64
import os
import re
import logging
import random
from datetime import date
from docx import Document
from docxtpl import DocxTemplate
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor, Pt, Inches
import xlsxwriter

logger = logging.getLogger(__name__)

OUTPUT_DIR = "generated_documents"
os.makedirs(OUTPUT_DIR, exist_ok=True)


def _extract_tax_rate(brand_dna: dict) -> float:
    calc = brand_dna.get("calculation_methods")
    if isinstance(calc, dict):
        for key in ("tax_rate", "vat_rate", "gst_rate"):
            rate = calc.get(key)
            if rate is not None:
                try:
                    return float(rate)
                except (ValueError, TypeError):
                    pass
    vat_status = brand_dna.get("vat_tax_status") or ""
    if isinstance(vat_status, str):
        match = re.search(r'(\d+(?:\.\d+)?)\s*%', vat_status)
        if match:
            return float(match.group(1))
    return 0.0


def _get_brand_color(brand_dna: dict) -> str:
    """Returns 6-char uppercase hex (no #) for the primary brand color."""
    color = (brand_dna.get("primary_color_hex") or "").strip().lstrip('#')
    if re.fullmatch(r'[0-9A-Fa-f]{6}', color):
        return color.upper()

    prefs = (brand_dna.get("layout_preferences") or "").lower()
    color_keywords = [
        ("navy", "1B3A5C"), ("dark blue", "1B3A5C"), ("royal blue", "2E5EAA"),
        ("teal", "1A7A8A"), ("dark green", "1B5E20"), ("green", "2E7D32"),
        ("blue", "2E5EAA"), ("red", "C62828"), ("maroon", "7B1521"),
        ("burgundy", "7B1521"), ("purple", "6A1B9A"), ("orange", "E65100"),
        ("charcoal", "37474F"), ("grey", "424242"), ("gray", "424242"),
        ("black", "212121"),
    ]
    for keyword, hex_val in color_keywords:
        if keyword in prefs:
            return hex_val
    return "1B3A5C"


def _hex_to_rgb(hex_color: str) -> tuple:
    h = hex_color.lstrip('#')
    return int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)


_CURRENCY_SYMBOLS = {
    "GBP": "£", "USD": "$", "EUR": "€", "AUD": "A$", "CAD": "C$",
    "NZD": "NZ$", "CHF": "Fr", "JPY": "¥", "CNY": "¥", "INR": "₹",
    "ZAR": "R", "SGD": "S$", "HKD": "HK$", "SEK": "kr", "NOK": "kr",
    "DKK": "kr", "MXN": "$", "BRL": "R$", "AED": "د.إ",
}


def _sym(currency_code: str) -> str:
    return _CURRENCY_SYMBOLS.get((currency_code or "").upper(), currency_code or "")


def _set_cell_bg(cell, hex_color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn('w:shd')):
        tcPr.remove(existing)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color.upper())
    tcPr.append(shd)


def _set_cell_borders(cell, color="DDDDDD", sz="4"):
    tcPr = cell._tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(existing)
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), sz)
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _remove_cell_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(existing)
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'auto')
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _add_rule(doc, color="CCCCCC", space_before=4, space_after=4):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def _spacer(doc, before=0, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    return p


class DocumentFactory:

    @staticmethod
    def generate_docx(quote_data: dict, brand_dna: dict, output_filename: str) -> dict:
        filepath = os.path.join(OUTPUT_DIR, output_filename)
        doc = Document()

        section = doc.sections[0]
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

        currency = quote_data.get("currency") or brand_dna.get("currency") or "GBP"
        sym = _sym(currency)
        brand_hex = _get_brand_color(brand_dna)
        brand_rgb = _hex_to_rgb(brand_hex)
        b_name = (brand_dna.get("business_name") or "Your Business").upper()
        today_str = date.today().strftime("%d %B %Y")
        quote_ref = f"QTE-{date.today().strftime('%Y%m')}-{random.randint(100, 999)}"
        logo_b64 = brand_dna.get("logo_base64")

        # ── 1. HEADER ─────────────────────────────────────────────────────
        info_lines = []
        if brand_dna.get("contact_details"):
            info_lines.append(str(brand_dna["contact_details"]))
        if brand_dna.get("vat_tax_status"):
            info_lines.append(str(brand_dna["vat_tax_status"]))
        if brand_dna.get("bank_info"):
            info_lines.append(f"Bank: {brand_dna['bank_info']}")

        if logo_b64:
            # Logo layout: logo image left | QUOTATION right (white bg), then full-width brand bar
            hdr_tbl = doc.add_table(rows=1, cols=2)
            hdr_tbl.autofit = False
            lc = hdr_tbl.cell(0, 0)
            rc = hdr_tbl.cell(0, 1)
            lc.width = Inches(3.5)
            rc.width = Inches(3.0)
            _remove_cell_borders(lc)
            _remove_cell_borders(rc)

            lp = lc.paragraphs[0]
            lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            lp.paragraph_format.space_before = Pt(4)
            lp.paragraph_format.space_after = Pt(4)
            try:
                img_stream = io.BytesIO(base64.b64decode(logo_b64))
                lp.add_run().add_picture(img_stream, height=Inches(0.55))
            except Exception as e:
                logger.warning(f"Could not insert logo image: {e}")
                lr = lp.add_run(b_name)
                lr.bold = True
                lr.font.size = Pt(14)

            rp = rc.paragraphs[0]
            rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            rp.paragraph_format.space_before = Pt(10)
            rp.paragraph_format.space_after = Pt(10)
            rr = rp.add_run("QUOTATION")
            rr.bold = True
            rr.font.size = Pt(20)

            # Full-width brand-coloured bar below logo/title
            bar_tbl = doc.add_table(rows=1, cols=1)
            bar_tbl.autofit = False
            bar_cell = bar_tbl.cell(0, 0)
            bar_cell.width = Inches(6.5)
            _set_cell_bg(bar_cell, brand_hex)
            _remove_cell_borders(bar_cell)
            bar_cell.paragraphs[0].paragraph_format.space_before = Pt(4)
            bar_cell.paragraphs[0].paragraph_format.space_after = Pt(4)

            # Company info below the bar
            if info_lines:
                ip = doc.add_paragraph()
                ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
                ip.paragraph_format.space_before = Pt(4)
                ip.paragraph_format.space_after = Pt(2)
                ip.add_run("\n".join(info_lines)).font.size = Pt(8.5)
        else:
            # Fallback: full-width brand-coloured band with business name + QUOTATION
            hdr_tbl = doc.add_table(rows=1, cols=2)
            hdr_tbl.autofit = False
            lc = hdr_tbl.cell(0, 0)
            rc = hdr_tbl.cell(0, 1)
            lc.width = Inches(4.0)
            rc.width = Inches(2.5)
            _set_cell_bg(lc, brand_hex)
            _set_cell_bg(rc, brand_hex)
            _remove_cell_borders(lc)
            _remove_cell_borders(rc)

            lp = lc.paragraphs[0]
            lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            lp.paragraph_format.space_before = Pt(8)
            lp.paragraph_format.space_after = Pt(8)
            lr = lp.add_run(b_name)
            lr.bold = True
            lr.font.color.rgb = RGBColor(255, 255, 255)
            lr.font.size = Pt(16)

            rp = rc.paragraphs[0]
            rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            rp.paragraph_format.space_before = Pt(8)
            rp.paragraph_format.space_after = Pt(8)
            rr = rp.add_run("QUOTATION")
            rr.bold = True
            rr.font.color.rgb = RGBColor(255, 255, 255)
            rr.font.size = Pt(16)

            if info_lines:
                ip = doc.add_paragraph()
                ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
                ip.paragraph_format.space_before = Pt(6)
                ip.paragraph_format.space_after = Pt(2)
                ip.add_run("\n".join(info_lines)).font.size = Pt(8.5)

        _add_rule(doc, color=brand_hex, space_before=4, space_after=4)

        # ── 3. QUOTE REF / DATE BAND ──────────────────────────────────────
        meta_tbl = doc.add_table(rows=1, cols=4)
        meta_tbl.autofit = False
        meta_entries = [
            ("Quote Ref:", True), (quote_ref, False),
            ("Quote Date:", True), (today_str, False),
        ]
        for i, (text, is_label) in enumerate(meta_entries):
            cell = meta_tbl.cell(0, i)
            cell.width = Inches(1.625)
            _set_cell_bg(cell, brand_hex if is_label else "F0F4F8")
            _remove_cell_borders(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(text)
            r.font.size = Pt(9)
            r.bold = is_label
            if is_label:
                r.font.color.rgb = RGBColor(255, 255, 255)

        _spacer(doc, before=0, after=10)

        # ── 4. BILL TO ────────────────────────────────────────────────────
        cust_name = quote_data.get("customer_name", "Customer")
        cust_addr = quote_data.get("customer_address")

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run("Bill To:")
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(*brand_rgb)

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(cust_name)
        r.bold = True
        r.font.size = Pt(10)

        if cust_addr:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.add_run(cust_addr).font.size = Pt(10)

        _spacer(doc, before=0, after=10)

        # ── 5. LINE ITEMS TABLE ───────────────────────────────────────────
        line_items = quote_data.get("line_items", [])
        subtotal = 0.0

        col_widths = [3.5, 0.6, 1.2, 1.2]
        col_aligns = [
            WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT,
        ]
        headers = ["Description", "Qty", f"Unit Price ({sym})", f"Total ({sym})"]

        tbl = doc.add_table(rows=1, cols=4)
        tbl.autofit = False

        hdr_cells = tbl.rows[0].cells
        for i, (h, align, w) in enumerate(zip(headers, col_aligns, col_widths)):
            hdr_cells[i].width = Inches(w)
            _set_cell_bg(hdr_cells[i], brand_hex)
            _set_cell_borders(hdr_cells[i], brand_hex)
            p = hdr_cells[i].paragraphs[0]
            p.alignment = align
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(h)
            r.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.font.size = Pt(9)

        for idx, item in enumerate(line_items):
            desc = str(item.get("description", ""))
            qty = float(item.get("quantity", 1))
            price = float(item.get("unit_price", 0.0))
            row_total = qty * price
            subtotal += row_total

            row = tbl.add_row()
            row_bg = "FFFFFF" if idx % 2 == 0 else "EEF2F7"
            vals = [desc, f"{qty:.0f}", f"{price:,.2f}", f"{row_total:,.2f}"]

            for i, (val, align, w) in enumerate(zip(vals, col_aligns, col_widths)):
                row.cells[i].width = Inches(w)
                _set_cell_bg(row.cells[i], row_bg)
                _set_cell_borders(row.cells[i])
                p = row.cells[i].paragraphs[0]
                p.alignment = align
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(3)
                p.add_run(val).font.size = Pt(9)

        # ── 6. TOTALS TABLE ───────────────────────────────────────────────
        tax_rate = _extract_tax_rate(brand_dna)
        tax_amount = subtotal * (tax_rate / 100)
        grand_total = subtotal + tax_amount

        _spacer(doc, before=0, after=2)

        totals_tbl = doc.add_table(rows=0, cols=2)
        totals_tbl.autofit = False

        def _add_total_row(label, value, bold=False, colored=False):
            row = totals_tbl.add_row()
            lc2 = row.cells[0]
            vc = row.cells[1]
            lc2.width = Inches(5.0)
            vc.width = Inches(1.5)
            _remove_cell_borders(lc2)
            _remove_cell_borders(vc)
            if colored:
                _set_cell_bg(lc2, brand_hex)
                _set_cell_bg(vc, brand_hex)

            lp2 = lc2.paragraphs[0]
            lp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            lp2.paragraph_format.space_before = Pt(2)
            lp2.paragraph_format.space_after = Pt(2)
            lr2 = lp2.add_run(label)
            lr2.bold = bold
            lr2.font.size = Pt(9)
            if colored:
                lr2.font.color.rgb = RGBColor(255, 255, 255)

            vp = vc.paragraphs[0]
            vp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            vp.paragraph_format.space_before = Pt(2)
            vp.paragraph_format.space_after = Pt(2)
            vr2 = vp.add_run(value)
            vr2.bold = bold
            vr2.font.size = Pt(9)
            if colored:
                vr2.font.color.rgb = RGBColor(255, 255, 255)

        _add_total_row("Subtotal:", f"{sym}{subtotal:,.2f}")
        if tax_rate > 0:
            _add_total_row(f"VAT/Tax ({tax_rate:.0f}%):", f"{sym}{tax_amount:,.2f}")
        _add_total_row("TOTAL DUE:", f"{sym}{grand_total:,.2f}", bold=True, colored=True)

        # ── 7. FOOTER ─────────────────────────────────────────────────────
        _spacer(doc, before=12, after=0)
        _add_rule(doc, color="CCCCCC", space_before=2, space_after=4)

        footer_parts = [b_name]
        if brand_dna.get("contact_details"):
            footer_parts.append(str(brand_dna["contact_details"]))
        if brand_dna.get("vat_tax_status"):
            vat = str(brand_dna["vat_tax_status"])
            _tax_kw = ("vat", "tax", "gst", "hst", "no tax")
            footer_parts.append(vat if any(k in vat.lower() for k in _tax_kw) else f"VAT No: {vat}")
        if brand_dna.get("bank_info"):
            footer_parts.append(str(brand_dna["bank_info"]))

        fp = doc.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.space_before = Pt(0)
        fr = fp.add_run(" | ".join(footer_parts))
        fr.font.size = Pt(7.5)
        fr.font.color.rgb = RGBColor(100, 100, 100)

        doc.save(filepath)
        logger.info(f"Generated DOCX: {filepath}")
        return {"filepath": filepath, "subtotal": subtotal, "tax_amount": tax_amount, "total": grand_total}

    @staticmethod
    def generate_xlsx(quote_data: dict, brand_dna: dict, output_filename: str) -> dict:
        filepath = os.path.join(OUTPUT_DIR, output_filename)
        currency = quote_data.get("currency") or brand_dna.get("currency") or "GBP"
        sym = _sym(currency)
        b_name = (brand_dna.get("business_name") or "Your Business").upper()
        brand_hex = _get_brand_color(brand_dna)
        brand_color_str = f"#{brand_hex}"
        today_str = date.today().strftime("%d %B %Y")
        quote_ref = f"QTE-{date.today().strftime('%Y%m')}-{random.randint(100, 999)}"

        workbook = xlsxwriter.Workbook(filepath)
        worksheet = workbook.add_worksheet("Quote")

        bold = workbook.add_format({'bold': True})
        header_fmt = workbook.add_format({
            'bold': True,
            'font_color': '#FFFFFF',
            'bg_color': brand_color_str,
            'border': 1,
            'border_color': brand_color_str,
            'align': 'center',
            'valign': 'vcenter',
        })
        header_left_fmt = workbook.add_format({
            'bold': True,
            'font_color': '#FFFFFF',
            'bg_color': brand_color_str,
            'border': 1,
            'border_color': brand_color_str,
            'valign': 'vcenter',
        })
        header_right_fmt = workbook.add_format({
            'bold': True,
            'font_color': '#FFFFFF',
            'bg_color': brand_color_str,
            'border': 1,
            'border_color': brand_color_str,
            'align': 'right',
            'valign': 'vcenter',
        })
        money_fmt = workbook.add_format({'num_format': '#,##0.00', 'border': 1, 'align': 'right'})
        money_alt_fmt = workbook.add_format({
            'num_format': '#,##0.00', 'border': 1, 'align': 'right',
            'bg_color': '#EEF2F7',
        })
        text_alt_fmt = workbook.add_format({'bg_color': '#EEF2F7', 'border': 1})
        text_border_fmt = workbook.add_format({'border': 1})
        qty_fmt = workbook.add_format({'border': 1, 'align': 'center'})
        qty_alt_fmt = workbook.add_format({'border': 1, 'align': 'center', 'bg_color': '#EEF2F7'})
        total_label_fmt = workbook.add_format({
            'bold': True, 'font_color': '#FFFFFF', 'bg_color': brand_color_str,
            'align': 'right', 'num_format': '#,##0.00',
        })
        total_val_fmt = workbook.add_format({
            'bold': True, 'font_color': '#FFFFFF', 'bg_color': brand_color_str,
            'align': 'right', 'num_format': '#,##0.00',
        })
        meta_label_fmt = workbook.add_format({
            'bold': True, 'font_color': '#FFFFFF', 'bg_color': brand_color_str,
            'align': 'center',
        })
        meta_val_fmt = workbook.add_format({
            'bg_color': '#F0F4F8', 'align': 'center',
        })
        small_grey = workbook.add_format({'font_size': 8, 'font_color': '#666666'})

        worksheet.set_column('A:A', 40)
        worksheet.set_column('B:B', 8)
        worksheet.set_column('C:C', 16)
        worksheet.set_column('D:D', 16)

        row = 0

        # Business name header
        worksheet.merge_range(row, 0, row, 2, b_name, workbook.add_format({
            'bold': True, 'font_size': 14,
            'font_color': '#FFFFFF', 'bg_color': brand_color_str,
            'valign': 'vcenter',
        }))
        worksheet.write(row, 3, "QUOTATION", workbook.add_format({
            'bold': True, 'font_size': 14,
            'font_color': '#FFFFFF', 'bg_color': brand_color_str,
            'align': 'right', 'valign': 'vcenter',
        }))
        worksheet.set_row(row, 24)
        row += 1

        # Company info
        info_lines = []
        if brand_dna.get("contact_details"):
            info_lines.append(str(brand_dna["contact_details"]))
        if brand_dna.get("vat_tax_status"):
            info_lines.append(str(brand_dna["vat_tax_status"]))
        if brand_dna.get("bank_info"):
            info_lines.append(f"Bank: {brand_dna['bank_info']}")
        for line in info_lines:
            worksheet.merge_range(row, 0, row, 3, line, small_grey)
            row += 1
        row += 1

        # Quote ref / date band
        worksheet.write(row, 0, "Quote Ref:", meta_label_fmt)
        worksheet.write(row, 1, quote_ref, meta_val_fmt)
        worksheet.write(row, 2, "Quote Date:", meta_label_fmt)
        worksheet.write(row, 3, today_str, meta_val_fmt)
        row += 2

        # Bill to
        worksheet.write(row, 0, "Bill To:", workbook.add_format({
            'bold': True, 'font_color': f'#{brand_hex}',
        }))
        row += 1
        worksheet.write(row, 0, quote_data.get("customer_name", "Customer"), bold)
        row += 1
        if quote_data.get("customer_address"):
            worksheet.write(row, 0, quote_data["customer_address"])
            row += 1
        row += 1

        # Table headers
        worksheet.write(row, 0, 'Description', header_left_fmt)
        worksheet.write(row, 1, 'Qty', header_fmt)
        worksheet.write(row, 2, f'Unit Price ({sym})', header_right_fmt)
        worksheet.write(row, 3, f'Total ({sym})', header_right_fmt)
        row += 1

        # Line items
        subtotal = 0.0
        for idx, item in enumerate(quote_data.get('line_items', [])):
            desc = item.get('description', '')
            qty = float(item.get('quantity', 1))
            price = float(item.get('unit_price', 0.0))
            total_line = qty * price
            subtotal += total_line
            alt = idx % 2 == 1
            worksheet.write(row, 0, desc, text_alt_fmt if alt else text_border_fmt)
            worksheet.write(row, 1, qty, qty_alt_fmt if alt else qty_fmt)
            worksheet.write_number(row, 2, price, money_alt_fmt if alt else money_fmt)
            worksheet.write_number(row, 3, total_line, money_alt_fmt if alt else money_fmt)
            row += 1

        row += 1

        # Totals
        tax_rate = _extract_tax_rate(brand_dna)
        tax_amount = subtotal * (tax_rate / 100)
        grand_total = subtotal + tax_amount

        subtotal_fmt = workbook.add_format({'num_format': '#,##0.00', 'align': 'right'})
        worksheet.write(row, 2, "Subtotal:", workbook.add_format({'bold': True, 'align': 'right'}))
        worksheet.write_number(row, 3, subtotal, subtotal_fmt)
        row += 1

        if tax_rate > 0:
            worksheet.write(row, 2, f"VAT/Tax ({tax_rate:.0f}%):", workbook.add_format({'bold': True, 'align': 'right'}))
            worksheet.write_number(row, 3, tax_amount, subtotal_fmt)
            row += 1

        worksheet.write(row, 2, "TOTAL DUE:", total_label_fmt)
        worksheet.write_number(row, 3, grand_total, total_val_fmt)
        row += 2

        # Footer
        footer_parts = [b_name]
        if brand_dna.get("contact_details"):
            footer_parts.append(str(brand_dna["contact_details"]))
        if brand_dna.get("vat_tax_status"):
            vat = str(brand_dna["vat_tax_status"])
            _tax_kw = ("vat", "tax", "gst", "hst", "no tax")
            footer_parts.append(vat if any(k in vat.lower() for k in _tax_kw) else f"VAT No: {vat}")
        if brand_dna.get("bank_info"):
            footer_parts.append(str(brand_dna["bank_info"]))
        worksheet.merge_range(row, 0, row, 3, " | ".join(footer_parts), small_grey)

        workbook.close()
        logger.info(f"Generated XLSX: {filepath}")
        return {"filepath": filepath, "subtotal": subtotal, "tax_amount": tax_amount, "total": grand_total}

    @staticmethod
    def generate_from_template(template_bytes: bytes, quote_data: dict, brand_dna: dict, output_filename: str) -> dict:
        """
        Renders a docxtpl quote template with actual quote data.
        Falls back to generate_docx if rendering fails.
        """
        filepath = os.path.join(OUTPUT_DIR, output_filename)
        currency = quote_data.get("currency") or brand_dna.get("currency") or "GBP"
        sym = _sym(currency)
        tax_rate = _extract_tax_rate(brand_dna)
        today_str = date.today().strftime("%d %B %Y")
        quote_ref = f"QTE-{date.today().strftime('%Y%m')}-{random.randint(100, 999)}"

        subtotal = sum(
            float(item.get("quantity", 1)) * float(item.get("unit_price", 0))
            for item in quote_data.get("line_items", [])
        )
        tax_amount = subtotal * (tax_rate / 100)
        grand_total = subtotal + tax_amount

        context = {
            "customer_name": quote_data.get("customer_name", ""),
            "customer_address": quote_data.get("customer_address") or "",
            "quote_ref": quote_ref,
            "quote_date": today_str,
            "subtotal": f"{sym}{subtotal:,.2f}",
            "tax_label": f"VAT/Tax ({tax_rate:.0f}%)" if tax_rate > 0 else "",
            "tax_amount": f"{sym}{tax_amount:,.2f}" if tax_rate > 0 else "",
            "grand_total": f"{sym}{grand_total:,.2f}",
            "line_items": [
                {
                    "description": str(item.get("description", "")),
                    "qty": f"{float(item.get('quantity', 1)):.0f}",
                    "unit_price_str": f"{sym}{float(item.get('unit_price', 0)):,.2f}",
                    "total_str": f"{sym}{float(item.get('quantity', 1)) * float(item.get('unit_price', 0)):,.2f}",
                }
                for item in quote_data.get("line_items", [])
            ],
        }

        try:
            tpl = DocxTemplate(io.BytesIO(template_bytes))
            tpl.render(context)
            tpl.save(filepath)
            logger.info(f"Generated DOCX from template: {filepath}")
            return {"filepath": filepath, "subtotal": subtotal, "tax_amount": tax_amount, "total": grand_total}
        except Exception as e:
            logger.error(f"Template rendering failed, falling back to scratch generation: {e}")
            return DocumentFactory.generate_docx(quote_data, brand_dna, output_filename)
