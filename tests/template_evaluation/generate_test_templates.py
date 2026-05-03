import os
import random
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_template(filename: str, template_type: str, include_logo: str = None):
    doc = Document()
    
    # 1. Handle Brand / Business Info based on type
    if include_logo and os.path.exists(include_logo):
        doc.add_picture(include_logo, width=Inches(1.5))
        
    p_biz = doc.add_paragraph()
    p_biz.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p_biz.add_run("ACME TRADES LTD\n123 Business Road\nLondon\nSW1A 1AA\ncontact@acmetrades.co.uk\n07000 123456")
    run.bold = True
    if template_type == "blue_theme":
        run.font.color.rgb = RGBColor(0, 51, 153) # Navy blue
    elif template_type == "orange_theme":
        run.font.color.rgb = RGBColor(255, 102, 0) # Orange
        
    doc.add_paragraph() # Spacing
    
    # 2. Document Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = p_title.add_run("QUOTATION")
    title_run.font.size = Pt(24)
    title_run.bold = True
    
    doc.add_paragraph()
    
    # 3. Customer Info Section (Varying layouts)
    layout = random.choice(["explicit_brackets", "blank_underscores", "implicit_cells", "no_address"])
    
    if layout == "explicit_brackets":
        doc.add_paragraph("PREPARED FOR:")
        doc.add_paragraph("Name: [Customer Name]")
        doc.add_paragraph("Address: [Customer Address]")
        doc.add_paragraph("Email: [Client Email]")
        doc.add_paragraph("Phone: [07xxx xxx xxx]")
    elif layout == "blank_underscores":
        doc.add_paragraph("Bill To:")
        doc.add_paragraph("Customer: _____________________")
        doc.add_paragraph("Address: _____________________")
        doc.add_paragraph("Email: _____________________")
    elif layout == "implicit_cells":
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Client:"
        table.cell(0, 1).text = "" # Should map to customer_name
        table.cell(1, 0).text = "Address:"
        table.cell(1, 1).text = "" # Should map to customer_address
    elif layout == "no_address":
        doc.add_paragraph("Client: [Name]")
        
    doc.add_paragraph()
    
    # 4. Meta Info (Quote Ref, Date, Expiry)
    meta_layout = random.choice(["inline_brackets", "table", "missing_expiry"])
    if meta_layout == "inline_brackets":
        doc.add_paragraph("Quote Ref: [Quote Reference]")
        doc.add_paragraph("Date: [Date]")
        doc.add_paragraph("Valid Until: [Expiry Date]")
    elif meta_layout == "table":
        table = doc.add_table(rows=3, cols=2)
        table.cell(0, 0).text = "Quote No"
        table.cell(0, 1).text = "[Ref]"
        table.cell(1, 0).text = "Issue Date"
        table.cell(1, 1).text = "[Date]"
        table.cell(2, 0).text = "Valid Until"
        table.cell(2, 1).text = "[Valid To]"
    elif meta_layout == "missing_expiry":
        doc.add_paragraph("Ref: _________")
        doc.add_paragraph("Date: ________")
        
    doc.add_paragraph()
        
    # 5. Line Items Table
    li_table = doc.add_table(rows=2, cols=4)
    li_table.style = 'Table Grid'
    hdr_cells = li_table.rows[0].cells
    hdr_cells[0].text = "Description of Work"
    hdr_cells[1].text = "Qty"
    hdr_cells[2].text = "Unit Price"
    hdr_cells[3].text = "Line Total"
    
    # Make header bold
    for cell in hdr_cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                
    # Add dummy row
    row_cells = li_table.rows[1].cells
    row_cells[0].text = "Sample item"
    row_cells[1].text = "1"
    row_cells[2].text = "100.00"
    row_cells[3].text = "100.00"
    
    doc.add_paragraph()
    
    # 6. Totals Section
    totals_layout = random.choice(["paragraph", "table"])
    if totals_layout == "paragraph":
        p_sub = doc.add_paragraph("Subtotal: ")
        p_sub.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_tax = doc.add_paragraph("VAT: ")
        p_tax.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_tot = doc.add_paragraph("Grand Total: ")
        p_tot.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_tot.runs[0].bold = True
    elif totals_layout == "table":
        # Totals in a separate table
        tot_table = doc.add_table(rows=3, cols=2)
        tot_table.alignment = 2 # Right align
        tot_table.cell(0, 0).text = "Sub Total"
        tot_table.cell(1, 0).text = "Tax"
        tot_table.cell(2, 0).text = "Total Due"
        
    # Add some footer text
    doc.add_paragraph()
    p_footer = doc.add_paragraph("Thank you for your business. Payment is due within 14 days.")
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save document
    doc.save(filename)
    print(f"Generated {filename} (Layout: {layout}, Meta: {meta_layout}, Totals: {totals_layout})")

if __name__ == "__main__":
    output_dir = "d:/Antigravity/Quote App/tests/template_evaluation"
    logos_dir = os.path.join(output_dir, "dummy_logos")
    
    # Get any logo files if they exist
    logo_files = []
    if os.path.exists(logos_dir):
        logo_files = [os.path.join(logos_dir, f) for f in os.listdir(logos_dir) if f.endswith(('.png', '.jpg'))]
    
    for i in range(1, 21):
        filename = os.path.join(output_dir, f"test_template_{i:02d}.docx")
        theme = random.choice(["blue_theme", "orange_theme", "default"])
        logo = random.choice(logo_files) if logo_files and random.random() > 0.5 else None
        
        create_template(filename, theme, logo)
    
    print("Done generating 20 templates.")
