import os
import sys
import io

# Add the parent directory to sys.path to import ai_service
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..')))

from ai_service import AIService
from docx import Document

def extract_jinja_tags(docx_bytes):
    doc = Document(io.BytesIO(docx_bytes))
    tags = set()
    
    # Check paragraphs
    for p in doc.paragraphs:
        txt = p.text
        if "{{" in txt or "{%" in txt:
            tags.add(txt.strip())
            
    # Check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    txt = p.text.strip()
                    if "{{" in txt or "{%" in txt:
                        tags.add(txt)
                        
    return tags

def run_evaluation():
    output_dir = "d:/Antigravity/Quote App/tests/template_evaluation"
    
    results = []
    
    for i in range(1, 21):
        filename = f"test_template_{i:02d}.docx"
        filepath = os.path.join(output_dir, filename)
        
        if not os.path.exists(filepath):
            continue
            
        print(f"\n--- Evaluating {filename} ---")
        try:
            # 1. Extract Brand DNA
            print("Extracting Brand DNA...")
            brand_dna = AIService.extract_brand_dna_from_blank(filepath)
            
            if not brand_dna:
                print("Failed to extract Brand DNA.")
                results.append({"file": filename, "success": False, "error": "Brand DNA extraction failed"})
                continue
                
            print(f"Brand DNA found: {brand_dna.get('business_name')} (Color: {brand_dna.get('primary_color_hex')}, Logo: {'Yes' if brand_dna.get('logo_base64') else 'No'})")
            
            # 2. Build template
            print("Building Quote Template...")
            template_bytes = AIService.build_quote_template(filepath, brand_dna)
            
            if not template_bytes:
                print("Failed to build quote template.")
                results.append({"file": filename, "success": False, "error": "Template building failed"})
                continue
                
            # 3. Analyze injected tags
            tags = extract_jinja_tags(template_bytes)
            print(f"Injected tags found: {tags}")
            
            # Save the processed template for manual inspection
            out_filepath = os.path.join(output_dir, f"processed_{filename}")
            with open(out_filepath, "wb") as f:
                f.write(template_bytes)
                
            results.append({
                "file": filename, 
                "success": True, 
                "tags": tags,
                "dna": brand_dna
            })
            
        except Exception as e:
            print(f"Error processing {filename}: {e}")
            results.append({"file": filename, "success": False, "error": str(e)})

    # Generate summary report
    print("\n\n================ EVALUATION SUMMARY ================")
    success_count = sum(1 for r in results if r["success"])
    print(f"Successfully processed: {success_count} / {len(results)}")
    
    # Expected fields
    standard_fields = ["{{ customer_name }}", "{{ quote_ref }}", "{{ quote_date }}", "{{ subtotal }}", "{{ grand_total }}"]
    
    for r in results:
        print(f"\n[{r['file']}]")
        if not r["success"]:
            print(f"  FAILED: {r.get('error')}")
            continue
            
        tags = r["tags"]
        # Convert tags to a flat string for easier matching
        tags_str = " ".join(tags)
        
        missing = [f for f in standard_fields if f not in tags_str]
        if missing:
            print(f"  WARNING: Missing expected tags: {missing}")
        else:
            print(f"  SUCCESS: Found all standard tags.")
            
        if "{%tr for item in line_items %}" not in tags_str:
            print("  WARNING: Line items loop NOT injected.")
        else:
            print("  SUCCESS: Line items loop injected.")

if __name__ == "__main__":
    run_evaluation()
