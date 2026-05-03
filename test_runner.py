"""
test_runner.py — Quote generation automated test suite.

Workflow:
  1. Creates 10 realistic DOCX templates (blank templates with bracket placeholders)
  2. Runs each through: extract_brand_dna → build_quote_template → generate_from_template
  3. Scores each on 5 dimensions (placeholder accuracy, line items, calculations,
     template fidelity, client readiness)
  4. Writes test_results/quote_test_analysis.md

Usage (from project root):  python test_runner.py
"""

import os
import re
import sys
import json
import random
import shutil
import logging
import argparse
import traceback
from pathlib import Path
from datetime import date, datetime

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.path.insert(0, str(Path(__file__).parent))
from ai_service import AIService
from document_factory import DocumentFactory, _sym
from config import settings
from google import genai

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s %(levelname)s %(message)s",
    datefmt="%H:%M:%S",
)
log = logging.getLogger("test_runner")

TEMPLATES_DIR = Path("test_templates")
GENERATED_DIR = Path("test_generated")
RESULTS_DIR = Path("test_results")
for _d in (TEMPLATES_DIR, GENERATED_DIR, RESULTS_DIR):
    _d.mkdir(exist_ok=True)

_gemini_client = genai.Client(api_key=settings.GEMINI_API_KEY)
ANALYSIS_MODEL = "gemini-2.5-flash"

# ── Field style helper ────────────────────────────────────────────────────────
# Controls how fillable fields appear in the blank template being tested:
#   "brackets" → [PlaceholderText]   (explicit markers, easiest for AI)
#   "labels"   → blank cell          (label column only, AI infers from context)
#   "sample"   → realistic text      (no markers, AI must detect from structure)
def _field_val(style: str, bracket_text: str, sample_text: str = "") -> str:
    if style == "labels":
        return ""
    if style == "sample":
        return sample_text
    return bracket_text  # default: "brackets"


# ─────────────────────────────────────────────────────────────────────────────
# TEMPLATE CONFIGURATIONS
# ─────────────────────────────────────────────────────────────────────────────

TEMPLATES = [
    {
        "id": "01_aura_design",
        "field_style": "brackets",
        "business_name": "Aura Design Studio",
        "business_address": "Studio 12, Neon Works, London, E1 6QL",
        "contact_line": "hello@auradesign.io  |  www.auradesign.io",
        "logo_path": "logos/aura_design.png",
        "columns": ["Project Item", "Amount"],
        "sample_rows": [
            ["Brand Identity System (Concept + Style Guide)", "£2,500.00"],
            ["Web UI Kit - 15 Responsive Screens", "£1,800.00"],
        ],
        "totals_rows": [("Total Investment", "£4,300.00")],
        "client_label": "Design Prepared For",
        "address_label": "Client Location",
        "ref_label": "Project ID",
        "date_label": "Issuance Date",
        "primary_color": "000000",
        "currency": "GBP",
        "tax_rate": 0.0,
        "quote_data": {
            "customer_name": "Zenith Ventures",
            "customer_address": "88 Kingsway, London, WC2B 6AA",
            "line_items": [
                {"description": "Brand Identity System (Concept + Style Guide)", "quantity": 1, "unit_price": 2500.00},
                {"description": "Web UI Kit - 15 Responsive Screens", "quantity": 1, "unit_price": 1800.00},
            ],
        },
    },
    {
        "id": "02_global_steel",
        "field_style": "labels",
        "business_name": "Global Steel & Supply Corp",
        "business_address": "1000 Industrial Way, Pittsburgh, PA 15201",
        "contact_line": "sales@globalsteel.com  |  Phone: 412-555-0100",
        "logo_path": "logos/global_steel.png",
        "vat_reg": "EIN: 12-3456789",
        "columns": ["Part No", "Description", "Qty", "Weight (lbs)", "Unit Price", "Subtotal"],
        "sample_rows": [
            ["ST-405", "Structural Steel Beam (H-Section) 20ft", "10", "4500", "$1,250.00", "$12,500.00"],
            ["RB-12", "Reinforcement Bars (Grade 60) 1/2in", "500", "3300", "$4.50", "$2,250.00"],
        ],
        "totals_rows": [("Subtotal", "$14,750.00"), ("Sales Tax (7%)", "$1,032.50"), ("TOTAL AMOUNT", "$15,782.50")],
        "client_label": "Consignee",
        "address_label": "Shipping Address",
        "ref_label": "PO REF NO",
        "date_label": "SHIP DATE",
        "primary_color": "D35400",
        "currency": "USD",
        "tax_rate": 7.0,
        "quote_data": {
            "customer_name": "Midwest Bridge Builders",
            "customer_address": "500 Construction Road, Chicago, IL 60601",
            "line_items": [
                {"description": "ST-405 Structural Steel Beam 20ft", "quantity": 10, "unit_price": 1250.00},
                {"description": "RB-12 Reinforcement Bars 1/2in", "quantity": 500, "unit_price": 4.50},
                {"description": "Heavy Duty Steel Coupling Pins", "quantity": 100, "unit_price": 12.00},
            ],
        },
    },
    {
        "id": "03_azure_estate",
        "field_style": "sample",
        "business_name": "The Azure Estate",
        "business_address": "Valley Road, Napa Valley, CA 94558",
        "contact_line": "weddings@azureestate.com  |  Tel: 707-555-8888",
        "logo_path": "logos/azure_estate.png",
        "columns": ["Service Category", "Description", "Investment"],
        "sample_rows": [
            ["Venue", "Exclusive use of the Grand Ballroom & Terrace", "$8,500.00"],
            ["Catering", "Premium 5-course dinner for 100 guests", "$12,000.00"],
        ],
        "totals_rows": [("Subtotal", "$20,500.00"), ("Service Charge (18%)", "$3,690.00"), ("GRAND TOTAL", "$24,190.00")],
        "client_label": "Honoring",
        "address_label": "Billing Contact",
        "ref_label": "Reservation #",
        "date_label": "Event Date",
        "primary_color": "2E4053",
        "currency": "USD",
        "tax_rate": 18.0,
        "quote_data": {
            "customer_name": "The Miller-Ross Wedding",
            "customer_address": "123 Vineyard Lane, St. Helena, CA 94574",
            "line_items": [
                {"description": "Exclusive use of the Grand Ballroom & Terrace", "quantity": 1, "unit_price": 8500.00},
                {"description": "Premium 5-course dinner per guest", "quantity": 100, "unit_price": 120.00},
                {"description": "Bridal Suite Accommodation (2 nights)", "quantity": 1, "unit_price": 1500.00},
            ],
        },
    },
    {
        "id": "04_quick_handyman",
        "field_style": "brackets",
        "business_name": "Quick Fix Handyman Services",
        "business_address": "12 Maple Ave, Springfield",
        "contact_line": "Call Mike: 555-0123",
        "columns": ["Job Description", "Price"],
        "sample_rows": [
            ["Repair leaky faucet in kitchen", "$85.00"],
            ["Replace door handle on front door", "$45.00"],
        ],
        "totals_rows": [("TOTAL DUE", "$130.00")],
        "client_label": "Customer",
        "address_label": "Job Site",
        "ref_label": "Job ID",
        "date_label": "Date",
        "primary_color": "2874A6",
        "currency": "USD",
        "tax_rate": 0.0,
        "quote_data": {
            "customer_name": "Alice Peterson",
            "customer_address": "456 Oak Lane, Springfield",
            "line_items": [
                {"description": "Repair leaky faucet in kitchen", "quantity": 1, "unit_price": 85.00},
                {"description": "Replace door handle on front door", "quantity": 1, "unit_price": 45.00},
                {"description": "Clean gutters (rear of house)", "quantity": 1, "unit_price": 120.00},
            ],
        },
    },
    {
        "id": "05_vet_clinic",
        "field_style": "labels",
        "business_name": "Happy Paws Veterinary Clinic",
        "business_address": "77 Bark Way, Melbourne VIC 3000",
        "contact_line": "vets@happypaws.com.au  |  (03) 9876 5432",
        "columns": ["Procedure / Item", "Qty", "Cost (inc GST)"],
        "sample_rows": [
            ["Annual Vaccination & Health Check", "1", "$110.00"],
            ["Dental Scale and Polish", "1", "$350.00"],
        ],
        "totals_rows": [("GST Included", "$41.82"), ("TOTAL PAYABLE", "$460.00")],
        "client_label": "Pet Owner",
        "address_label": "Residential Address",
        "ref_label": "Patient Record #",
        "date_label": "Visit Date",
        "primary_color": "148F77",
        "currency": "AUD",
        "tax_rate": 10.0,
        "quote_data": {
            "customer_name": "Sarah Miller (Pet: Luna)",
            "customer_address": "15 Wattle St, Richmond VIC 3121",
            "line_items": [
                {"description": "Annual Vaccination & Health Check", "quantity": 1, "unit_price": 110.00},
                {"description": "Dental Scale and Polish", "quantity": 1, "unit_price": 350.00},
                {"description": "Antibiotic Course (10 days)", "quantity": 1, "unit_price": 45.00},
            ],
        },
    },
    {
        "id": "06_solar_rebate",
        "field_style": "sample",
        "business_name": "Eco-Power Solar Solutions",
        "business_address": "Unit 5, Energy Park, Sydney NSW 2000",
        "contact_line": "info@ecopower.com.au  |  ABN: 99 888 777 666",
        "columns": ["System Component", "Qty", "Price"],
        "sample_rows": [
            ["6.6kW Tier 1 Solar Panel Array", "1", "$6,500.00"],
            ["5kW Hybrid Inverter", "1", "$2,200.00"],
        ],
        "totals_rows": [
            ("Gross System Total", "$8,700.00"),
            ("Govt Rebate (STCs)", "-$2,400.00"),
            ("NET AMOUNT PAYABLE", "$6,300.00")
        ],
        "client_label": "Applicant Name",
        "address_label": "Installation Site",
        "ref_label": "System Quote Ref",
        "date_label": "Expiry Date",
        "primary_color": "F1C40F",
        "currency": "AUD",
        "tax_rate": 0.0,
        "quote_data": {
            "customer_name": "Thomas Jenkins",
            "customer_address": "82 Botany Rd, Alexandria NSW 2015",
            "line_items": [
                {"description": "6.6kW Tier 1 Solar Panel Array", "quantity": 1, "unit_price": 6500.00},
                {"description": "5kW Hybrid Inverter (Battery Ready)", "quantity": 1, "unit_price": 2200.00},
                {"description": "Smart Meter Installation", "quantity": 1, "unit_price": 450.00},
            ],
        },
    },
    {
        "id": "07_move_easy",
        "field_style": "labels",
        "business_name": "Move Easy Logisitics",
        "business_address": "202 Cargo Way, Dallas, TX 75201",
        "contact_line": "ops@moveeasy.com  |  DOT: 1234567",
        "columns": ["Service Description", "Volume (cu.ft)", "Rate/cu.ft", "Total"],
        "sample_rows": [
            ["Residential Move - 3 Bedroom House", "1200", "$4.50", "$5,400.00"],
            ["Packing & Unpacking Services", "1", "flat", "$850.00"],
        ],
        "totals_rows": [("Fuel Surcharge (10%)", "$625.00"), ("ESTIMATED TOTAL", "$6,875.00")],
        "client_label": "Customer Name",
        "address_label": "Origin Address",
        "ref_label": "Booking Ref",
        "date_label": "Proposed Move Date",
        "primary_color": "C0392B",
        "currency": "USD",
        "tax_rate": 0.0,
        "quote_data": {
            "customer_name": "Robert Henderson",
            "customer_address": "77 Sunset Blvd, Dallas, TX 75204",
            "line_items": [
                {"description": "Residential Move - 3 Bedroom House", "quantity": 1200, "unit_price": 4.50},
                {"description": "Packing & Unpacking Services (Flat Rate)", "quantity": 1, "unit_price": 850.00},
                {"description": "Full Insurance Coverage (Up to $50k)", "quantity": 1, "unit_price": 250.00},
            ],
        },
    },
    {
        "id": "08_code_crafters",
        "field_style": "brackets",
        "business_name": "Code Crafters Software",
        "business_address": "Tech Plaza, Berlin, Germany",
        "contact_line": "billing@codecrafters.de  |  VAT: DE987654321",
        "columns": ["Software Service", "Seats", "Monthly Fee", "Yearly Total"],
        "sample_rows": [
            ["Enterprise License - CRM Suite", "25", "€45.00", "€13,500.00"],
            ["Premium Support Add-on", "1", "€200.00", "€2,400.00"],
        ],
        "totals_rows": [("Subtotal (Annual)", "€15,900.00"), ("VAT (19%)", "€3,021.00"), ("TOTAL COST", "€18,921.00")],
        "client_label": "Account Owner",
        "address_label": "Entity Address",
        "ref_label": "License ID",
        "date_label": "Renewal Date",
        "primary_color": "2471A3",
        "currency": "EUR",
        "tax_rate": 19.0,
        "quote_data": {
            "customer_name": "Innova Solutions GmbH",
            "customer_address": "Alte Strasse 12, 10115 Berlin",
            "line_items": [
                {"description": "Enterprise License - CRM Suite", "quantity": 25, "unit_price": 540.00}, # 45 * 12
                {"description": "Premium Support Add-on (Annual)", "quantity": 1, "unit_price": 2400.00},
            ],
        },
    },
    {
        "id": "09_gourmet_catering",
        "field_style": "sample",
        "business_name": "Gourmet Garden Catering",
        "business_address": "High Street, Oxford, OX1 4AH",
        "contact_line": "events@gourmetgarden.co.uk",
        "columns": ["Item Description", "Price Per Head", "Quantity", "Total"],
        "sample_rows": [
            ["Artisan Hors d'oeuvres Selection", "£12.50", "120", "£1,500.00"],
            ["Plated Main Course - Seasonal Menu", "£45.00", "120", "£5,400.00"],
        ],
        "totals_rows": [("Subtotal", "£6,900.00"), ("Staffing Fee", "£850.00"), ("TOTAL QUOTE", "£7,750.00")],
        "client_label": "Event Organized For",
        "address_label": "Venue Address",
        "ref_label": "Event Reference",
        "date_label": "Booking Date",
        "primary_color": "1D8348",
        "currency": "GBP",
        "tax_rate": 0.0,
        "quote_data": {
            "customer_name": "Oxford Alumni Association",
            "customer_address": "Sheldonian Theatre, Oxford, OX1 3AZ",
            "line_items": [
                {"description": "Artisan Hors d'oeuvres Selection", "quantity": 120, "unit_price": 12.50},
                {"description": "Plated Main Course - Seasonal Menu", "quantity": 120, "unit_price": 45.00},
                {"description": "Dessert Buffet & Coffee Station", "quantity": 120, "unit_price": 8.50},
            ],
        },
    },
    {
        "id": "10_fit_pro",
        "field_style": "labels",
        "business_name": "FitPro Personal Training",
        "business_address": "Health Hub, Gym Street, Manchester",
        "contact_line": "coach@fitpro.com",
        "columns": ["Training Package", "Cost"],
        "sample_rows": [
            ["10 Session Kickstart Pack", "£450.00"],
            ["Nutritional Consultation", "£85.00"],
        ],
        "totals_rows": [("TOTAL", "£535.00")],
        "client_label": "Client",
        "address_label": "Contact Info",
        "ref_label": "Membership #",
        "date_label": "Start Date",
        "primary_color": "7B241C",
        "currency": "GBP",
        "tax_rate": 0.0,
        "quote_data": {
            "customer_name": "Mark Stevens",
            "customer_address": "12 Broadway, Manchester, M1 1AA",
            "line_items": [
                {"description": "10 Session Kickstart Pack", "quantity": 1, "unit_price": 450.00},
                {"description": "Nutritional Consultation", "quantity": 1, "unit_price": 85.00},
                {"description": "FitPro T-Shirt & Water Bottle", "quantity": 1, "unit_price": 0.00},
            ],
        },
    },
]


# Synthetic quote data pool — cycled over when running against real user templates.
_SYNTH_QUOTE_POOL = [t["quote_data"] for t in TEMPLATES]
_SYNTH_CURRENCY_POOL = [(t["currency"], t["tax_rate"]) for t in TEMPLATES]


def load_existing_templates(max_count: int = 10) -> list[dict]:
    """
    Fetch real user blank templates from Supabase Storage.
    Returns a list of minimal config dicts ready for run_pipeline().
    """
    from supabase import create_client as _sb_create

    log.info("Connecting to Supabase to list user templates...")
    try:
        sb = _sb_create(settings.SUPABASE_URL, settings.SUPABASE_KEY)
        user_dirs = sb.storage.from_(settings.SUPABASE_TEMPLATES_BUCKET).list("templates")
    except Exception as e:
        log.error(f"Failed to list Supabase templates bucket: {e}")
        return []

    # Each entry is a folder named by user UUID
    user_ids = [item["name"] for item in user_dirs if item.get("name") and "." not in item["name"]]
    if not user_ids:
        log.warning("No user template folders found in Supabase.")
        return []

    random.shuffle(user_ids)
    log.info(f"Found {len(user_ids)} user template folder(s) — sampling up to {max_count}.")

    templates = []
    for user_id in user_ids:
        if len(templates) >= max_count:
            break
        storage_path = f"templates/{user_id}/blank_template.docx"
        try:
            data = sb.storage.from_(settings.SUPABASE_TEMPLATES_BUCKET).download(storage_path)
        except Exception:
            log.warning(f"No blank_template.docx for user {user_id} — skipping.")
            continue

        n = len(templates) + 1
        tid = f"existing_{n:02d}"
        local_path = TEMPLATES_DIR / f"{tid}_blank.docx"
        local_path.write_bytes(data)
        log.info(f"Downloaded: {storage_path} → {local_path}")

        # Borrow currency/tax/quote_data from the synthetic pool (cycling)
        idx = (n - 1) % len(_SYNTH_QUOTE_POOL)
        currency, tax_rate = _SYNTH_CURRENCY_POOL[idx]
        quote_data = _SYNTH_QUOTE_POOL[idx]

        templates.append({
            "id": tid,
            "business_name": None,   # unknown — skip in fidelity scoring
            "business_address": None,
            "columns": ["(unknown)"],
            "currency": currency,
            "tax_rate": tax_rate,
            "quote_data": quote_data,
            "blank_path_override": local_path,
        })

    log.info(f"Loaded {len(templates)} existing template(s) from Supabase.")
    return templates


# ─────────────────────────────────────────────────────────────────────────────
# DOCX TEMPLATE CREATION HELPERS
# ─────────────────────────────────────────────────────────────────────────────

def _rgb(hex_color: str) -> RGBColor:
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _shade_cell(cell, fill_hex: str):
    tcPr = cell._tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn("w:shd")):
        tcPr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex.upper().lstrip("#"))
    tcPr.append(shd)


def _set_cell_text_styled(cell, text: str, bold: bool = False, color_hex: str = None, size_pt: int = None):
    para = cell.paragraphs[0]
    for run in para.runs:
        run.text = ""
    if para.runs:
        run = para.runs[0]
        run.text = text
    else:
        run = para.add_run(text)
    if bold:
        run.font.bold = True
    if color_hex:
        run.font.color.rgb = _rgb(color_hex)
    if size_pt:
        run.font.size = Pt(size_pt)


def create_template_doc(t: dict) -> Document:
    """Creates a DOCX blank template for the given config dict."""
    doc = Document()
    primary = t["primary_color"]
    white = "FFFFFF"
    tid = t["id"]
    fs = t.get("field_style", "brackets")

    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    # ── Logo support ─────────────────────────────────────────────────────────
    logo_path = t.get("logo_path")
    if logo_path and os.path.exists(logo_path):
        # Center logo for Aura Design, Top Left for others
        if tid == "01_aura_design":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run()
            r.add_picture(logo_path, width=Inches(1.5))
        else:
            p = doc.add_paragraph()
            r = p.add_run()
            r.add_picture(logo_path, width=Inches(1.2))
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # ── Business header ──────────────────────────────────────────────────────
    # For Aura Design, center everything and make it minimalist
    if tid == "01_aura_design":
        header_p = doc.add_paragraph()
        header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = header_p.add_run(t["business_name"].upper())
        r.font.size = Pt(14)
        r.font.bold = True
        r.font.color.rgb = _rgb(primary)
        
        addr_p = doc.add_paragraph(f"{t['business_address']} | {t['contact_line']}")
        addr_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if addr_p.runs:
            addr_p.runs[0].font.size = Pt(9)
        doc.add_paragraph("").paragraph_format.space_after = Pt(20)
    else:
        p = doc.add_paragraph()
        r = p.add_run(t["business_name"])
        r.font.size = Pt(16)
        r.font.bold = True
        r.font.color.rgb = _rgb(primary)
        p.paragraph_format.space_after = Pt(2)

        addr_p = doc.add_paragraph(t["business_address"])
        addr_p.paragraph_format.space_before = Pt(0)
        addr_p.paragraph_format.space_after = Pt(0)

        contact_p = doc.add_paragraph(t["contact_line"])
        contact_p.paragraph_format.space_before = Pt(0)
        contact_p.paragraph_format.space_after = Pt(0)

        if t.get("vat_reg"):
            vp = doc.add_paragraph(t["vat_reg"])
            vp.paragraph_format.space_before = Pt(0)
            vp.paragraph_format.space_after = Pt(4)

    # ── QUOTATION title ──────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    if tid == "01_aura_design":
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        title_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
    tr = title_p.add_run("QUOTATION")
    tr.font.size = Pt(18)
    tr.font.bold = True
    tr.font.color.rgb = _rgb(primary)
    title_p.paragraph_format.space_after = Pt(8)

    # ── Meta table: ref + date ───────────────────────────────────────────────
    ref_label = t["ref_label"]
    date_label = t["date_label"]
    
    # Global Steel wants a more complex meta section
    if tid == "02_global_steel":
        meta = doc.add_table(rows=3, cols=4)
    else:
        meta = doc.add_table(rows=2, cols=4)
        
    meta.style = "Table Grid"
    label_bg = "F4F6F7"

    meta.cell(0, 0).text = ref_label
    _shade_cell(meta.cell(0, 0), label_bg)
    meta.cell(0, 1).text = _field_val(fs, f"[{ref_label}]", "QT-1001")

    meta.cell(0, 2).text = date_label
    _shade_cell(meta.cell(0, 2), label_bg)
    meta.cell(0, 3).text = _field_val(fs, f"[{date_label}]", "15 Jan 2024")

    meta.cell(1, 0).text = "Valid Until"
    _shade_cell(meta.cell(1, 0), label_bg)
    meta.cell(1, 1).text = _field_val(fs, "[EXPIRY_DATE]", "14 Feb 2024")

    if tid == "02_global_steel":
        meta.cell(2, 0).text = "Project Code"
        _shade_cell(meta.cell(2, 0), label_bg)
        meta.cell(2, 1).text = _field_val(fs, "[PROJECT_REF]", "PRJ-001")
        meta.cell(2, 2).text = "Department"
        _shade_cell(meta.cell(2, 2), label_bg)
        meta.cell(2, 3).text = _field_val(fs, "[DEPT_NAME]", "Procurement")
    else:
        meta.cell(1, 2).text = "Status"
        _shade_cell(meta.cell(1, 2), label_bg)
        meta.cell(1, 3).text = _field_val(fs, "[QUOTE_STATUS]", "Draft")

    doc.add_paragraph("").paragraph_format.space_before = Pt(6)

    # ── Bill To ──────────────────────────────────────────────────────────────
    bt = doc.add_paragraph()
    label_text = "BILL TO:" if tid != "03_azure_estate" else "HONORING:"
    bt_run = bt.add_run(label_text)
    bt_run.font.bold = True
    bt_run.font.color.rgb = _rgb(primary)
    bt.paragraph_format.space_after = Pt(0)

    doc.add_paragraph(_field_val(fs, f"[{t['client_label']}]", "ABC Company Ltd")).paragraph_format.space_after = Pt(0)
    doc.add_paragraph(_field_val(fs, f"[{t['address_label']}]", "100 Example Street, City")).paragraph_format.space_after = Pt(8)

    doc.add_paragraph("")

    # ── Line items table ─────────────────────────────────────────────────────
    columns = t["columns"]
    sample_rows = t["sample_rows"]
    num_cols = len(columns)
    items_tbl = doc.add_table(rows=1 + len(sample_rows), cols=num_cols)
    items_tbl.style = "Table Grid"

    for ci, col_name in enumerate(columns):
        cell = items_tbl.cell(0, ci)
        _set_cell_text_styled(cell, col_name, bold=True, color_hex=white)
        _shade_cell(cell, primary)

    for ri, row_data in enumerate(sample_rows):
        for ci, val in enumerate(row_data):
            items_tbl.cell(ri + 1, ci).text = str(val)

    doc.add_paragraph("").paragraph_format.space_before = Pt(4)

    # ── Totals table ─────────────────────────────────────────────────────────
    totals_rows = t["totals_rows"]
    totals_tbl = doc.add_table(rows=len(totals_rows), cols=2)
    # Right align the totals table
    totals_tbl.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    totals_tbl.style = "Table Grid"

    for ri, (label, val) in enumerate(totals_rows):
        label_cell = totals_tbl.cell(ri, 0)
        val_cell = totals_tbl.cell(ri, 1)
        is_last = ri == len(totals_rows) - 1
        if is_last:
            _shade_cell(label_cell, primary)
            _shade_cell(val_cell, primary)
            _set_cell_text_styled(label_cell, label, bold=True, color_hex=white)
            _set_cell_text_styled(val_cell, val, bold=True, color_hex=white)
        else:
            label_cell.text = label
            val_cell.text = val

    doc.add_paragraph("").paragraph_format.space_before = Pt(8)

    # ── Terms / Footer ────────────────────────────────────────────────────────
    pmt = doc.add_paragraph()
    pr = pmt.add_run("TERMS & CONDITIONS:")
    pr.font.bold = True
    pr.font.color.rgb = _rgb(primary)
    pmt.paragraph_format.space_after = Pt(0)

    doc.add_paragraph("Please review the above quote. Valid for 14 days.").paragraph_format.space_after = Pt(0)
    if tid == "08_code_crafters":
        doc.add_paragraph("Annual billing applies. Early cancellation fees may incur.").paragraph_format.space_after = Pt(0)

    return doc



# ─────────────────────────────────────────────────────────────────────────────
# PIPELINE RUNNER
# ─────────────────────────────────────────────────────────────────────────────

def run_pipeline(t: dict) -> dict:
    """
    Runs one template through the full pipeline.
    Returns a result dict with paths, brand_dna, quote_data, error info.
    """
    tid = t["id"]
    result = {
        "id": tid,
        "config": t,
        "blank_path": None,
        "processed_path": None,
        "generated_path": None,
        "brand_dna": None,
        "errors": [],
        "warnings": [],
    }

    # Step 1: Create blank template (skipped for existing/real-user templates)
    blank_path = t.get("blank_path_override") or TEMPLATES_DIR / f"{tid}_blank.docx"
    if t.get("blank_path_override"):
        log.info(f"[{tid}] Using pre-downloaded blank template: {blank_path}")
        result["blank_path"] = blank_path
    else:
        log.info(f"[{tid}] Creating blank template...")
        try:
            doc = create_template_doc(t)
            doc.save(str(blank_path))
            result["blank_path"] = blank_path
            log.info(f"[{tid}] Blank template saved: {blank_path}")
        except Exception as e:
            result["errors"].append(f"Template creation failed: {e}")
            log.error(f"[{tid}] Template creation error: {e}")
            return result

    # Step 2: Extract brand DNA
    log.info(f"[{tid}] Extracting brand DNA...")
    try:
        brand_dna = AIService.extract_brand_dna_from_blank(str(blank_path))
        if not brand_dna:
            result["errors"].append("extract_brand_dna_from_blank returned None")
            return result
        # Merge in currency/tax/format (normally collected during bot onboarding)
        brand_dna["currency"] = t["currency"]
        brand_dna["calculation_methods"] = {"tax_rate": t["tax_rate"]}
        brand_dna["preferred_format"] = "docx"
        result["brand_dna"] = brand_dna
        log.info(f"[{tid}] Brand DNA extracted: business_name={brand_dna.get('business_name')!r}, color={brand_dna.get('primary_color_hex')!r}")
    except Exception as e:
        result["errors"].append(f"Brand DNA extraction failed: {e}")
        log.error(f"[{tid}] Brand DNA error: {traceback.format_exc()}")
        return result

    # Step 3: Build Jinja2 template
    log.info(f"[{tid}] Building quote template (Gemini field mapping)...")
    try:
        template_bytes = AIService.build_quote_template(str(blank_path), brand_dna)
        if not template_bytes:
            result["errors"].append("build_quote_template returned None")
            return result
        processed_path = TEMPLATES_DIR / f"{tid}_processed.docx"
        processed_path.write_bytes(template_bytes)
        result["processed_path"] = processed_path
        log.info(f"[{tid}] Processed template saved: {processed_path}")
    except Exception as e:
        result["errors"].append(f"Template building failed: {e}")
        log.error(f"[{tid}] build_quote_template error: {traceback.format_exc()}")
        return result

    # Step 4: Generate quote
    log.info(f"[{tid}] Generating quote...")
    output_filename = f"{tid}_generated.docx"
    try:
        gen_result = DocumentFactory.generate_from_template(
            template_bytes,
            t["quote_data"],
            brand_dna,
            output_filename,
        )
        src = Path(gen_result["filepath"])
        dst = GENERATED_DIR / f"{tid}_generated.docx"
        shutil.copy2(str(src), str(dst))
        result["generated_path"] = dst
        result["gen_financials"] = {
            "subtotal": gen_result["subtotal"],
            "tax_amount": gen_result["tax_amount"],
            "total": gen_result["total"],
        }
        log.info(f"[{tid}] Quote generated: {dst}  subtotal={gen_result['subtotal']:.2f}")
    except Exception as e:
        result["errors"].append(f"Quote generation failed: {e}")
        log.error(f"[{tid}] generate_from_template error: {traceback.format_exc()}")

    return result


# ─────────────────────────────────────────────────────────────────────────────
# ANALYSIS
# ─────────────────────────────────────────────────────────────────────────────

def extract_docx_text(path: Path) -> str:
    """Extract all text from a DOCX file."""
    try:
        doc = Document(str(path))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())
        return "\n".join(parts)
    except Exception as e:
        return f"[TEXT EXTRACTION FAILED: {e}]"


_JINJA_RE = re.compile(r"\{\{.*?\}\}|\{%.*?%\}")
_BRACKET_RE = re.compile(r"\[.+?\]")


def score_placeholder_accuracy(text: str) -> tuple[int, list[str]]:
    """Score 0-10: check for unfilled Jinja2 or bracket placeholders."""
    issues = []
    jinja_hits = _JINJA_RE.findall(text)
    bracket_hits = [m for m in _BRACKET_RE.findall(text) if len(m) > 3]

    for hit in jinja_hits:
        issues.append(f"Unfilled Jinja2 tag: `{hit}`")
    for hit in bracket_hits:
        issues.append(f"Remaining bracket placeholder: `{hit}`")

    total_issues = len(jinja_hits) + len(bracket_hits)
    score = max(0, 10 - total_issues * 2)
    return score, issues


def score_line_items(text: str, quote_data: dict, sym: str) -> tuple[int, list[str]]:
    """Score 0-10: check all expected line items appear in the generated text."""
    issues = []
    line_items = quote_data.get("line_items", [])
    found = 0

    for item in line_items:
        desc = item.get("description", "")
        # Check for at least the first 20 chars of description
        needle = desc[:20].lower()
        if needle and needle in text.lower():
            found += 1
        else:
            issues.append(f"Line item description not found: '{desc[:40]}'")

    if not line_items:
        return 5, ["No line items in test data"]

    ratio = found / len(line_items)
    score = round(ratio * 10)
    if found == len(line_items):
        issues = []
    return score, issues


def score_calculations(
    text: str, quote_data: dict, tax_rate: float, sym: str
) -> tuple[int, list[str]]:
    """Score 0-10: verify subtotal, tax, and total appear correctly in the text."""
    issues = []
    items = quote_data.get("line_items", [])
    subtotal = sum(float(i.get("quantity", 1)) * float(i.get("unit_price", 0)) for i in items)
    tax_amount = subtotal * (tax_rate / 100)
    grand_total = subtotal + tax_amount

    def _fmt(amount: float) -> str:
        return f"{sym}{amount:,.2f}"

    checks = [
        (_fmt(subtotal), "subtotal"),
        (_fmt(grand_total), "grand total"),
    ]
    if tax_rate > 0:
        checks.append((_fmt(tax_amount), "tax amount"))

    found = 0
    for val_str, label in checks:
        # Strip currency symbol for a fallback match (handles multi-char symbols like A$)
        bare = val_str.lstrip("A$C$NZ$HK$S$R$£€¥₹")
        if val_str in text or bare in text:
            found += 1
        else:
            issues.append(f"Expected {label} {val_str!r} not found in output")

    score = round((found / len(checks)) * 10)
    return score, issues


def score_template_fidelity(text: str, t: dict) -> tuple[int, list[str]]:
    """Score 0-10: verify template business identity is preserved in the output."""
    issues = []
    found = 0
    total_checks = 1  # always check customer name

    # Business identity checks — skipped for real-user templates where name is unknown
    if t.get("business_name"):
        total_checks += 2
        for needle, label in [
            (t["business_name"], "business name"),
            (t.get("business_address", "").split(",")[0], "business address (first part)"),
        ]:
            if needle and needle.lower() in text.lower():
                found += 1
            else:
                issues.append(f"Template identity field not found — {label}: '{needle}'")

    # Check customer name was injected
    cust = t["quote_data"]["customer_name"]
    if cust.lower() in text.lower():
        found += 1
    else:
        issues.append(f"Customer name not injected: '{cust}'")

    score = round((found / total_checks) * 10)
    return score, issues


def score_client_readiness_gemini(
    blank_text: str, generated_text: str, template_id: str
) -> tuple[int, str]:
    """Score 0-10 via Gemini qualitative review of the generated quote."""
    prompt = f"""You are reviewing a generated business quote document to assess whether it is ready to send to a client.

Template ID: {template_id}

--- Original blank template (business identity only) ---
{blank_text[:1500]}

--- Generated quote output ---
{generated_text[:2500]}

Assess the generated quote on:
1. Professional appearance (readable, complete sentences, no raw placeholders)
2. Completeness (header, line items, totals, payment details all present)
3. Accuracy (customer name present, figures look plausible)
4. Client-readiness (could this be sent to a real client without embarrassment?)

Respond ONLY with valid JSON:
{{"score": <0-10>, "reasoning": "<2-3 sentence summary>", "issues": ["<issue1>", "<issue2>"]}}

Return an empty "issues" array if none found. Do not wrap in code fences."""

    try:
        from google.genai import types as _types
        resp = _gemini_client.models.generate_content(
            model=ANALYSIS_MODEL,
            contents=prompt,
            config=_types.GenerateContentConfig(response_mime_type="application/json"),
        )
        raw = resp.text or ""
        if raw.startswith("```"):
            raw = re.sub(r"^```[a-z]*\n?", "", raw).rstrip("```").strip()
        data = json.loads(raw)
        return int(data.get("score", 5)), data.get("reasoning", ""), data.get("issues", [])
    except Exception as e:
        log.warning(f"Gemini client readiness scoring failed for {template_id}: {e}")
        return 5, f"Scoring error: {e}", []


def analyze_result(result: dict) -> dict:
    """Run all scoring dimensions on a completed pipeline result."""
    t = result["config"]
    sym = _sym(t["currency"])
    tax_rate = t["tax_rate"]

    analysis = {
        "id": result["id"],
        "errors": result.get("errors", []),
        "warnings": result.get("warnings", []),
        "scores": {},
        "issues_by_dim": {},
    }

    if result.get("generated_path"):
        gen_text = extract_docx_text(result["generated_path"])
    else:
        analysis["scores"] = {d: 0 for d in ["placeholder", "line_items", "calculations", "fidelity", "client_readiness"]}
        analysis["issues_by_dim"]["pipeline"] = result.get("errors", ["Pipeline did not complete"])
        return analysis

    blank_text = extract_docx_text(result["blank_path"]) if result.get("blank_path") else ""

    s1, i1 = score_placeholder_accuracy(gen_text)
    s2, i2 = score_line_items(gen_text, t["quote_data"], sym)
    s3, i3 = score_calculations(gen_text, t["quote_data"], tax_rate, sym)
    s4, i4 = score_template_fidelity(gen_text, t)
    s5, reasoning5, i5 = score_client_readiness_gemini(blank_text, gen_text, t["id"])

    analysis["scores"] = {
        "placeholder_accuracy": s1,
        "line_items_integrity": s2,
        "calculation_accuracy": s3,
        "template_fidelity": s4,
        "client_readiness": s5,
        "total": s1 + s2 + s3 + s4 + s5,
    }
    analysis["issues_by_dim"] = {
        "placeholder_accuracy": i1,
        "line_items_integrity": i2,
        "calculation_accuracy": i3,
        "template_fidelity": i4,
        "client_readiness": i5,
    }
    analysis["client_readiness_reasoning"] = reasoning5
    return analysis


# ─────────────────────────────────────────────────────────────────────────────
# REPORT GENERATOR
# ─────────────────────────────────────────────────────────────────────────────

def write_report(results: list[dict], analyses: list[dict], run_date: str):
    """Write test_results/quote_test_analysis.md"""
    report_path = RESULTS_DIR / "quote_test_analysis.md"

    dim_keys = ["placeholder_accuracy", "line_items_integrity", "calculation_accuracy", "template_fidelity", "client_readiness"]
    dim_labels = {
        "placeholder_accuracy": "Placeholder Accuracy",
        "line_items_integrity": "Line Items Integrity",
        "calculation_accuracy": "Calculation Accuracy",
        "template_fidelity": "Template Fidelity",
        "client_readiness": "Client Readiness",
    }

    lines = []
    lines.append(f"# Quote Generation Test Results — {run_date}\n")
    lines.append(f"**Templates tested:** {len(results)}  |  **Run date:** {run_date}\n")

    # ── Summary table ────────────────────────────────────────────────────────
    lines.append("## Summary\n")
    header = "| Template | " + " | ".join(dim_labels[k][:4] + "." for k in dim_keys) + " | **Total /50** | Status |"
    sep = "|---|" + "---|" * len(dim_keys) + "---|---|"
    lines.append(header)
    lines.append(sep)

    all_issues = []

    for analysis in analyses:
        tid = analysis["id"]
        scores = analysis.get("scores", {})
        total = scores.get("total", 0)
        dims = " | ".join(str(scores.get(k, "–")) for k in dim_keys)
        status = "✅ PASS" if total >= 35 else ("⚠️ WARN" if total >= 20 else "❌ FAIL")
        if analysis.get("errors"):
            status = "💥 ERROR"
            total = 0
        lines.append(f"| `{tid}` | {dims} | **{total}** | {status} |")

        for dim, issues in analysis.get("issues_by_dim", {}).items():
            for issue in issues:
                all_issues.append((tid, dim_labels.get(dim, dim), issue))

    lines.append("")

    # ── Dimension descriptions ────────────────────────────────────────────────
    lines.append("### Scoring Dimensions\n")
    lines.append("| Dimension | What is measured |")
    lines.append("|---|---|")
    lines.append("| Placeholder Accuracy | No `{{ }}` Jinja2 or `[bracket]` placeholders remain |")
    lines.append("| Line Items Integrity | All requested line item descriptions appear in the output |")
    lines.append("| Calculation Accuracy | Subtotal, tax, and total figures are mathematically correct |")
    lines.append("| Template Fidelity | Business identity (name, address) and customer name are present |")
    lines.append("| Client Readiness | Gemini qualitative review — is this ready to send? |")
    lines.append("")

    # ── Per-template sections ─────────────────────────────────────────────────
    lines.append("## Per-Template Detail\n")

    for result, analysis in zip(results, analyses):
        t = result["config"]
        tid = t["id"]
        scores = analysis.get("scores", {})
        total = scores.get("total", 0)

        biz = t.get("business_name") or "(real user template)"
        cols = t.get("columns", ["(unknown)"])
        col_str = "(detected by AI)" if cols == ["(unknown)"] else ", ".join(cols)
        fstyle = t.get("field_style", "brackets")
        lines.append(f"### {tid}")
        lines.append(f"**Business:** {biz}  |  **Currency:** {t['currency']}  |  **Tax:** {t['tax_rate']}%  |  **Field Style:** `{fstyle}`  |  **Columns:** {col_str}\n")

        if analysis.get("errors"):
            lines.append(f"**Pipeline Errors:**")
            for e in analysis["errors"]:
                lines.append(f"- ❌ {e}")
            lines.append("")
            continue

        lines.append(f"**Total Score: {total}/50**\n")
        lines.append("| Dimension | Score | Issues |")
        lines.append("|---|---|---|")

        for k in dim_keys:
            s = scores.get(k, "–")
            issues = analysis.get("issues_by_dim", {}).get(k, [])
            issue_str = "; ".join(issues[:2]) if issues else "None"
            if len(issues) > 2:
                issue_str += f" (+{len(issues)-2} more)"
            lines.append(f"| {dim_labels[k]} | {s}/10 | {issue_str} |")

        if analysis.get("client_readiness_reasoning"):
            lines.append(f"\n**Client Readiness Note:** {analysis['client_readiness_reasoning']}")

        lines.append("")

    # ── Aggregated issues ─────────────────────────────────────────────────────
    lines.append("## All Issues Found\n")
    if all_issues:
        lines.append("| Template | Dimension | Issue |")
        lines.append("|---|---|---|")
        for tid, dim, issue in all_issues:
            lines.append(f"| `{tid}` | {dim} | {issue} |")
    else:
        lines.append("No issues found.")
    lines.append("")

    # ── Recommended fixes ─────────────────────────────────────────────────────
    lines.append("## Recommended Fixes & Next Steps\n")

    # Aggregate issues by dimension
    dim_issue_map: dict[str, list[tuple[str, str]]] = {}
    for tid, dim, issue in all_issues:
        dim_issue_map.setdefault(dim, []).append((tid, issue))

    if not dim_issue_map:
        lines.append("All tests passed with no issues. No changes recommended at this time.")
    else:
        for dim, items in dim_issue_map.items():
            lines.append(f"### {dim}")
            seen = set()
            for tid, issue in items:
                if issue not in seen:
                    lines.append(f"- **[{tid}]** {issue}")
                    seen.add(issue)
            lines.append("")

    lines.append("---")
    lines.append(f"*Generated by `test_runner.py` on {run_date}*")

    report_path.write_text("\n".join(lines), encoding="utf-8")
    log.info(f"Report written: {report_path}")
    return report_path


# ─────────────────────────────────────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="Quote generation test suite")
    parser.add_argument(
        "--mode",
        choices=["new", "existing"],
        default="new",
        help=(
            "new: create 10 synthetic DOCX templates from config (default); "
            "existing: fetch up to 10 real user templates from Supabase Storage"
        ),
    )
    args, _ = parser.parse_known_args()

    run_date = datetime.now().strftime("%Y-%m-%d %H:%M")
    log.info(f"=== Quote Generation Test Suite — {run_date} (mode: {args.mode}) ===")
    log.info(f"Templates: {TEMPLATES_DIR}  |  Generated: {GENERATED_DIR}  |  Results: {RESULTS_DIR}")

    if args.mode == "existing":
        templates_to_run = load_existing_templates(max_count=10)
        if not templates_to_run:
            log.error("No existing templates found — aborting. Try --mode new instead.")
            sys.exit(1)
    else:
        templates_to_run = TEMPLATES

    results = []
    analyses = []

    for i, t in enumerate(templates_to_run, 1):
        log.info(f"\n{'='*60}")
        biz = t.get("business_name") or "(real user template)"
        log.info(f"[{i}/{len(templates_to_run)}] Running: {t['id']} ({biz})")
        log.info(f"{'='*60}")
        result = run_pipeline(t)
        results.append(result)

    log.info(f"\n{'='*60}")
    log.info("Pipeline complete. Running analysis...")
    log.info(f"{'='*60}\n")

    for result in results:
        log.info(f"Analysing: {result['id']}...")
        analysis = analyze_result(result)
        analyses.append(analysis)
        total = analysis.get("scores", {}).get("total", "n/a")
        errors = len(analysis.get("errors", []))
        log.info(f"  → Score: {total}/50  |  Errors: {errors}")

    report_path = write_report(results, analyses, run_date)

    # Print summary to console
    print(f"\n{'='*60}")
    print(f"TEST RESULTS SUMMARY")
    print(f"{'='*60}")
    print(f"{'Template':<25} {'Total':>6}  Status")
    print(f"{'-'*45}")
    for analysis in analyses:
        tid = analysis["id"]
        total = analysis.get("scores", {}).get("total", 0)
        if analysis.get("errors"):
            status = "ERROR"
        elif total >= 35:
            status = "PASS"
        elif total >= 20:
            status = "WARN"
        else:
            status = "FAIL"
        print(f"{tid:<25} {total:>5}/50  {status}")
    print(f"\nReport: {report_path.absolute()}")


if __name__ == "__main__":
    main()
