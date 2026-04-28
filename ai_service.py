from google import genai
from google.genai import types
import asyncio
import io
import json
import logging
import random
import re
import time
import os
import base64 as _b64
from config import settings

logger = logging.getLogger(__name__)

try:
    import sentry_sdk as _sentry
except ImportError:
    _sentry = None


def _capture(e: Exception):
    if _sentry and settings.SENTRY_DSN:
        _sentry.capture_exception(e)

client = genai.Client(api_key=settings.GEMINI_API_KEY) if settings.GEMINI_API_KEY else None
if not client:
    logger.warning("GEMINI_API_KEY not found in environment.")

DNA_EXTRACTION_PROMPT = """
You are a highly capable AI assistant helping a tradesperson set up their automated quoting system.
Analyze the provided blank quote template and extract the business's "Brand DNA" into a strict JSON format.
The template contains the business's branding and structure but no client or job details.

Extract the following information:
1. "business_name": The name of the business issuing the quotes.
2. "business_address": The physical address of the business.
3. "contact_details": Phone numbers, emails, or websites.
4. "bank_info": Bank details for payment (Account Number, Sort Code, IBAN, etc.).
5. "layout_preferences": Notes on visual layout (e.g., "Logo top left", "Clean modern font", "Blue color scheme").
6. "primary_color_hex": The dominant brand/header color as a 6-digit hex code WITHOUT the # prefix (e.g., "1B3A5C" for navy blue). Look at header backgrounds, table header rows, colored banners, and logo colors. This is critical — inspect the document carefully for any colored elements.
7. "secondary_color_hex": The secondary accent color as a 6-digit hex code WITHOUT the # prefix, if present (e.g., highlight colors, subheadings). Set to null if none.

Return ONLY a valid JSON object with these keys. If a field cannot be determined, set its value to null.
"""

QUOTE_GENERATION_PROMPT = """
You are an expert AI assistant who turns raw user requests into structured quote data for tradespeople.
Input is often terse shorthand like "Customer Name. Job description Total" or "Name, job, price".
Extract the following from the user's input in strict JSON format:
1. "customer_name": Name of the customer (string)
2. "customer_address": Address of the customer if mentioned (string or null)
3. "line_items": An array of objects, each with:
   - "description" (string): what the work or material is
   - "quantity" (number): how many units
   - "unit_price" (number): price per unit

Rules:
- A trailing number after the job description is ALWAYS the price — never an address, postcode, or ID.
- If a lump-sum total is given (e.g. "New bathroom 3000"), create ONE line item with quantity=1 and unit_price equal to that total.
- If prices are not mentioned at all, infer reasonable defaults for the trade described.
- Always produce at least one line item — never return an empty line_items array.
- Fix any obvious spelling mistakes in the description fields (e.g. "bathrrom" → "bathroom").
- Capitalise the first letter of each description (e.g. "new bathroom" → "New bathroom").

Example:
Input: "Amy Smith. New bathroom 3000"
Output: {"customer_name": "Amy Smith", "customer_address": null, "line_items": [{"description": "New bathroom", "quantity": 1, "unit_price": 3000}]}

Return ONLY valid JSON.
"""

VOICE_QUOTE_PROMPT = """
You are an expert AI assistant who turns voice notes from tradespeople into structured quote data.
Transcribe the following voice note and extract the quote details in strict JSON format:
1. "customer_name": Name of the customer (string or null)
2. "customer_address": Address of the customer if mentioned (string or null)
3. "line_items": An array of objects, each with:
   - "description" (string): what the work or material is
   - "quantity" (number): how many units
   - "unit_price" (number): price per unit

Rules:
- If a lump-sum total is mentioned (e.g. "New bathroom, three thousand pounds"), create ONE line item with quantity=1 and unit_price equal to that total.
- If prices are not mentioned at all, infer reasonable defaults for the trade described.
- Always produce at least one line item — never return an empty line_items array.
- Fix any obvious spelling mistakes in the description fields.
- Capitalise the first letter of each description.
Return ONLY valid JSON.
"""

IMAGE_QUOTE_PROMPT = """
You are an expert AI assistant who extracts quote information from images sent by tradespeople.
The image may show handwritten notes, a photo of a job, a printed list, a whiteboard, or a notebook.
Extract the following in strict JSON format:
1. "customer_name": Name of the customer if visible (string or null)
2. "customer_address": Address if visible (string or null)
3. "line_items": An array of objects, each with:
   - "description" (string): what the work or material is
   - "quantity" (number): how many units
   - "unit_price" (number): price per unit

If prices are not visible, infer reasonable defaults based on the trade context.
Return ONLY valid JSON.
"""

REFINEMENT_PROMPT = """
You are helping a tradesperson refine a quote. They were shown a summary and have replied.

Current quote data (JSON):
{current_quote}

User's reply: "{user_response}"

Determine whether the user is CONFIRMING the quote (e.g. yes, ok, looks good, send it, correct, fine, that's right, generate it)
or REQUESTING CHANGES (e.g. mentioning different prices, names, items to add/remove, corrections, currency changes).

If the user requests a currency change (e.g. "currency is £", "use GBP", "change to euros", "currency should be USD"):
- Update the "currency" field in updated_quote to the correct ISO code (e.g. "GBP" for £/pounds, "EUR" for euros, "USD" for dollars).

Return a JSON object with exactly these two keys:
- "confirmed": true if confirming, false if requesting changes
- "updated_quote": the complete updated quote JSON (unchanged if confirmed, with modifications applied if changes were requested)
"""

TEMPLATE_MAP_PROMPT = """
You are analyzing a DOCX quote/invoice document to identify which cells contain VARIABLE data that changes per quote.

Document structure (JSON):
{structure}

Business info already extracted (static — do NOT mark these as variable):
{business_info}

Identify the locations of these VARIABLE fields. Return null for any field you cannot locate.

Note: tables are indexed in document order (0-based), including nested tables (e.g. inner tables used for side-by-side layouts). Outer layout/container cells will appear empty — focus on inner cells with specific field text.

Respond with ONLY a valid JSON object with these keys:
- "customer_name": location where the customer/client name appears
- "customer_address": location of customer address (null if absent)
- "quote_ref": location of quote/invoice reference number (null if absent)
- "quote_date": location of the date (null if absent)
- "line_items_table_index": integer index (0-based) of the table containing line items rows (description, qty, price, total columns)
- "line_item_header_rows": list of row indices (0-based) that are header rows in the line items table
- "subtotal_value_location": location of the subtotal numeric value cell (null if absent)
- "tax_amount_value_location": location of the tax/VAT amount numeric value cell (null if absent)
- "grand_total_value_location": location of the grand total / total due numeric value cell

Location format: {{"type": "table", "table_index": N, "row_index": N, "col_index": N}}
  OR {{"type": "paragraph", "paragraph_index": N}} (use 0-based index into the non-empty paragraphs list)
"""

XLSX_DNA_PROMPT = """
You are a highly capable AI assistant helping a tradesperson set up their automated quoting system.
Analyze the provided blank Excel quote template (cell values listed below) and extract the business's "Brand DNA" into a strict JSON format.
The template contains the business's branding and structure but no client or job details.

Extract the following information:
1. "business_name": The name of the business issuing the quotes.
2. "business_address": The physical address of the business.
3. "contact_details": Phone numbers, emails, or websites.
4. "bank_info": Bank details for payment (Account Number, Sort Code, IBAN, etc.).
5. "layout_preferences": Notes on visual layout (e.g., "Header row at top", "Blue color scheme").
6. "primary_color_hex": If you can infer a brand color from the template description, return it as a 6-digit hex WITHOUT the # prefix. Otherwise set to null.
7. "secondary_color_hex": Secondary accent color as a 6-digit hex WITHOUT the # prefix, or null if not determinable.

Return ONLY a valid JSON object with these keys. If a field cannot be determined, set its value to null.
"""

XLSX_MAP_PROMPT = """
You are analyzing an Excel quote/invoice template to identify where variable data should be written.
Below is a dump of the template's cell values (format: SheetName!CellRef: 'value').

Template cell values:
{cell_dump}

Business info already known (these are STATIC — do NOT mark as variable fields):
{business_info}

Identify the exact cell addresses for the following VARIABLE fields that change per quote.
Return null for any field you cannot locate.

For line items, identify:
- The 1-based row number where the first data row begins (after headers)
- The column letters for description, qty, unit_price, and total columns

Respond with ONLY a valid JSON object with exactly these keys:
- "client_name": cell address (e.g. "B5") or null
- "client_address": cell address or null
- "quote_ref": cell address for quote/invoice reference number or null
- "quote_date": cell address for the date or null
- "line_items_start_row": integer (1-based row number of first data row) or null
- "line_items_cols": object with keys "description", "qty", "unit_price", "total" — each value is a column letter (e.g. "A") or null
- "subtotal_cell": cell address for subtotal value or null
- "tax_cell": cell address for tax/VAT amount or null
- "total_cell": cell address for grand total or null
"""

MODEL = "gemini-2.5-flash"
FALLBACK_MODEL = "gemini-2.5-flash-lite"
MAX_RETRIES = 5
BASE_RETRY_DELAY = 3  # seconds

JSON_CONFIG = types.GenerateContentConfig(response_mime_type="application/json")

# Limits concurrent Gemini calls to prevent thread pool exhaustion under load.
# Tune this based on your Gemini API quota (free tier: lower; paid tier: higher).
_ai_semaphore = asyncio.Semaphore(5)


async def run_ai(func, *args):
    """Run a sync AI function with concurrency limiting (max 5 simultaneous Gemini calls)."""
    async with _ai_semaphore:
        return await asyncio.to_thread(func, *args)


class RateLimitError(Exception):
    """Raised when Gemini API rate limit retries are exhausted."""
    pass


def _normalize_quote(quote: dict) -> dict:
    """Capitalise first letter of each line item description (client-side safety net)."""
    for item in quote.get("line_items", []):
        desc = item.get("description", "")
        if desc:
            item["description"] = desc[0].upper() + desc[1:]
    return quote


def _is_transient(error_str: str) -> bool:
    return "429" in error_str or "quota" in error_str.lower() or "503" in error_str or "UNAVAILABLE" in error_str


def _generate_with_retry(contents, config=JSON_CONFIG):
    """Calls Gemini with exponential backoff + jitter, falling back to lite model if primary exhausts."""
    for attempt in range(MAX_RETRIES):
        try:
            return client.models.generate_content(model=MODEL, contents=contents, config=config)
        except Exception as e:
            if _is_transient(str(e)):
                wait_time = BASE_RETRY_DELAY * (2 ** attempt) + random.uniform(0, 2)
                logger.warning(f"Gemini {MODEL} transient error (attempt {attempt + 1}/{MAX_RETRIES}), retrying in {wait_time:.1f}s: {e}")
                time.sleep(wait_time)
            else:
                _capture(e)
                raise

    logger.warning(f"{MODEL} exhausted after {MAX_RETRIES} attempts, trying fallback {FALLBACK_MODEL}")
    for attempt in range(2):
        try:
            return client.models.generate_content(model=FALLBACK_MODEL, contents=contents, config=config)
        except Exception as e:
            if _is_transient(str(e)):
                wait_time = BASE_RETRY_DELAY * (2 ** attempt) + random.uniform(0, 2)
                logger.warning(f"Gemini {FALLBACK_MODEL} transient error (attempt {attempt + 1}/2), retrying in {wait_time:.1f}s: {e}")
                time.sleep(wait_time)
            else:
                _capture(e)
                raise

    raise RateLimitError("Gemini API is unavailable — please try again in a moment.")


def _detect_line_items_table(all_tables_in_doc) -> int | None:
    """Fallback: score tables by header keywords to find the line items table."""
    desc_kw = {
        'description', 'item', 'service', 'work', 'details', 'particulars', 'goods', 'labour', 'material',
        'package', 'scope', 'task', 'deliverable', 'product', 'items', 'services', 'works', 'activity',
    }
    qty_kw = {'qty', 'quantity', 'units', 'hours', 'hrs', 'count', 'no.'}
    price_kw = {'price', 'rate', 'cost', 'unit price', 'unit rate', 'charge', 'fee', 'each'}
    total_kw = {'total', 'amount', 'sum', 'net', 'gross', 'line total', 'ex. vat', 'inc. vat'}

    best_idx, best_score = None, 0
    for idx, table in enumerate(all_tables_in_doc):
        if not table.rows or len(table.rows[0].cells) < 2:
            continue
        score = 0
        for cell in table.rows[0].cells:
            txt = cell.text.strip().lower()
            if any(k in txt for k in desc_kw):
                score += 3
            if any(k in txt for k in qty_kw):
                score += 2
            if any(k in txt for k in price_kw):
                score += 2
            if any(k in txt for k in total_kw):
                score += 1
        if score > best_score:
            best_score, best_idx = score, idx
    return best_idx


def _map_line_item_columns(header_cells) -> dict:
    """Map column indices to field names based on header text. Returns {col_idx: field_name}."""
    desc_kw = {
        'description', 'item', 'service', 'work', 'details', 'particulars', 'goods', 'labour', 'material',
        'package', 'scope', 'task', 'deliverable', 'product', 'items', 'services', 'works', 'activity',
    }
    qty_kw = {'qty', 'quantity', 'units', 'hours', 'hrs', 'count', 'no.'}
    # 'unit' alone (without 'price'/'cost') → unit-of-measure column
    unit_kw = {'unit', 'uom', 'measure'}
    price_kw = {'unit price', 'unit rate', 'rate', 'per unit', 'unit cost', 'each', 'price', 'fee', 'charge'}
    total_kw = {'total', 'amount', 'sum', 'net', 'gross', 'line total', 'ex. vat', 'inc. vat'}

    col_map: dict[int, str] = {}
    for ci, cell in enumerate(header_cells):
        txt = cell.text.strip().lower()
        if any(k in txt for k in desc_kw):
            col_map[ci] = 'description'
        elif any(k in txt for k in qty_kw):
            col_map[ci] = 'qty'
        elif any(k in txt for k in price_kw):
            col_map[ci] = 'unit_price'
        elif any(k in txt for k in total_kw):
            col_map[ci] = 'total'
        elif any(k in txt for k in unit_kw):
            col_map[ci] = 'unit'
    return col_map


class AIService:
    @staticmethod
    def extract_brand_dna_from_blank(docx_path: str) -> dict | None:
        """
        Takes a single blank DOCX template path, extracts brand identity (name,
        address, colors, logo) and returns the structured JSON Brand DNA dict,
        or None on failure. Currency and tax are NOT extracted here — they are
        collected via direct Q&A during onboarding.
        """
        try:
            import base64 as _b64
            import docx as _docx
            from docx.oxml.ns import qn as _qn

            logo_b64 = None
            docx_primary_color = None
            _raster_exts = {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.tif'}

            try:
                doc = _docx.Document(docx_path)
            except Exception as e:
                logger.error(f"Failed to open DOCX for brand DNA extraction: {e}")
                return None

            full_text = []

            # Regular paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)

            # Table cells
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text)

            # Text boxes — often used for company name/address
            try:
                for txbx in doc.element.body.iter(_qn('w:txbxContent')):
                    for t_el in txbx.iter(_qn('w:t')):
                        if t_el.text and t_el.text.strip():
                            full_text.append(t_el.text)
            except Exception:
                pass

            # Headers and footers — often contain business name/address
            try:
                for section in doc.sections:
                    for hdr_ftr in (section.header, section.footer):
                        if getattr(hdr_ftr, 'is_linked_to_previous', True):
                            continue
                        for para in hdr_ftr.paragraphs:
                            if para.text.strip():
                                full_text.append(para.text)
                        for tbl in hdr_ftr.tables:
                            for row in tbl.rows:
                                for cell in row.cells:
                                    if cell.text.strip():
                                        full_text.append(cell.text)
            except Exception:
                pass

            extracted = "\n".join(full_text)
            logger.info(f"DOCX text extracted: {len(extracted)} chars")

            if len(extracted) < 50:
                logger.warning("DOCX text extraction yielded very little content — template may use unsupported layout")

            # Extract primary brand color directly from DOCX XML (overrides Gemini's guess)
            try:
                from collections import Counter as _Counter
                import re as _re
                _IGNORE_COLORS = {
                    'FFFFFF', 'F0F4F8', 'EEF2F7', 'F5F5F5', 'EEEEEE',
                    'E0E0E0', 'DDDDDD', 'CCCCCC', 'AUTO', 'NONE', '',
                }
                color_counts = _Counter()
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            tcPr = cell._tc.find(_qn('w:tcPr'))
                            if tcPr is not None:
                                shd = tcPr.find(_qn('w:shd'))
                                if shd is not None:
                                    fill = shd.get(_qn('w:fill'), '').upper().strip()
                                    if fill not in _IGNORE_COLORS and _re.fullmatch(r'[0-9A-F]{6}', fill):
                                        color_counts[fill] += 1
                if color_counts:
                    docx_primary_color = color_counts.most_common(1)[0][0]
                    logger.info(f"DOCX primary color extracted from XML: {docx_primary_color}")
            except Exception as _ce:
                logger.warning(f"DOCX color extraction failed (non-fatal): {_ce}")

            # Extract first raster image as logo
            for rel in doc.part.rels.values():
                try:
                    target = getattr(rel, 'target_ref', '') or ''
                    ext = os.path.splitext(target.lower())[1]
                    if ext in _raster_exts:
                        blob = rel.target_part.blob
                        if len(blob) > 500:
                            logo_b64 = _b64.b64encode(blob).decode('utf-8')
                            break
                except Exception:
                    continue

            combined_text = DNA_EXTRACTION_PROMPT + f"\n\n--- Blank template content ---\n{extracted}"
            logger.info(f"Calling Gemini for brand DNA extraction, content length: {len(combined_text)} chars")
            response = _generate_with_retry(combined_text)

            try:
                raw = response.text or ""
            except Exception as e:
                logger.error(f"response.text raised: {e}", exc_info=True)
                return None

            logger.info(f"Brand DNA raw response (first 300 chars): {raw[:300]!r}")

            # Strip markdown code fences if Gemini wraps the JSON
            if raw.startswith("```"):
                raw = raw.split("```")[1]
                if raw.startswith("json"):
                    raw = raw[4:]
            raw = raw.strip()
            if not raw:
                logger.error("Gemini returned empty response for brand DNA extraction")
                return None

            result = json.loads(raw)
            if not isinstance(result, dict) or not result:
                logger.error(f"Gemini returned non-dict or empty result: {result!r}")
                return None

            # DOCX-extracted values take priority over Gemini's text analysis
            if logo_b64:
                result["logo_base64"] = logo_b64
            if docx_primary_color:
                result["primary_color_hex"] = docx_primary_color
            return result

        except RateLimitError:
            raise
        except Exception as e:
            logger.error(f"Failed to extract brand DNA: {e}", exc_info=True)
            _capture(e)
            return None

    @staticmethod
    def generate_quote_data(text: str) -> dict:
        """Parses user text into structured quote data using Gemini."""
        try:
            response = _generate_with_retry(f"{QUOTE_GENERATION_PROMPT}\n\nUser Input: {text}")
            return _normalize_quote(json.loads(response.text))
        except RateLimitError:
            raise
        except Exception as e:
            logger.error(f"Failed to generate quote data: {e}")
            _capture(e)
            return {}

    @staticmethod
    def transcribe_and_extract_voice(file_path: str) -> dict:
        """Transcribes a voice note (OGG) and extracts structured quote data."""
        try:
            uploaded_file = client.files.upload(
                file=file_path,
                config=types.UploadFileConfig(mime_type="audio/ogg"),
            )

            # Wait for Gemini to finish processing the uploaded file
            max_wait = 30
            waited = 0
            while getattr(uploaded_file.state, "name", str(uploaded_file.state)) == "PROCESSING":
                time.sleep(2)
                waited += 2
                uploaded_file = client.files.get(name=uploaded_file.name)
                if waited >= max_wait:
                    raise Exception("Timed out waiting for voice file to be processed by Gemini")

            if getattr(uploaded_file.state, "name", str(uploaded_file.state)) == "FAILED":
                raise Exception(f"Gemini file processing failed for {file_path}")

            response = _generate_with_retry([VOICE_QUOTE_PROMPT, uploaded_file])

            try:
                client.files.delete(name=uploaded_file.name)
            except Exception:
                pass

            return _normalize_quote(json.loads(response.text))

        except RateLimitError:
            raise
        except Exception as e:
            logger.error(f"Failed to process voice note: {e}")
            _capture(e)
            return {}

    @staticmethod
    def extract_quote_from_image(file_path: str) -> dict:
        """Extracts structured quote data from an image (photo of notes, job site, etc.)."""
        try:
            uploaded_file = client.files.upload(file=file_path)

            response = _generate_with_retry([IMAGE_QUOTE_PROMPT, uploaded_file])

            try:
                client.files.delete(name=uploaded_file.name)
            except Exception:
                pass

            return _normalize_quote(json.loads(response.text))

        except RateLimitError:
            raise
        except Exception as e:
            logger.error(f"Failed to extract quote from image: {e}")
            _capture(e)
            return {}

    @staticmethod
    def refine_quote(current_quote: dict, user_response: str) -> dict:
        """
        Determines if the user is confirming a quote or requesting changes.
        Returns {"confirmed": bool, "updated_quote": dict}.
        """
        try:
            prompt = REFINEMENT_PROMPT.format(
                current_quote=json.dumps(current_quote, indent=2),
                user_response=user_response
            )
            response = _generate_with_retry(prompt)
            result = json.loads(response.text)
            if "updated_quote" in result:
                result["updated_quote"] = _normalize_quote(result["updated_quote"])
            return result

        except RateLimitError:
            raise
        except Exception as e:
            logger.error(f"Failed to refine quote: {e}")
            _capture(e)
            return {"confirmed": True, "updated_quote": current_quote}

    @staticmethod
    def build_quote_template(docx_path: str, brand_dna: dict) -> bytes | None:
        """
        Takes a sample quote DOCX, asks Gemini to map variable fields, injects
        docxtpl Jinja2 placeholders via python-docx, and returns the modified
        DOCX as bytes. Returns None if the document cannot be processed.
        """
        try:
            import docx as _docx
            from docx.oxml.ns import qn
            from docx.table import Table as _DocxTable

            doc = _docx.Document(docx_path)
        except Exception as e:
            logger.error(f"Failed to open DOCX for template building: {e}")
            return None

        # ── Helpers ─────────────────────────────────────────────────────────

        def _set_para_text(para, text: str):
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = text
            else:
                para.add_run(text)

        def _set_cell_text(cell, text: str):
            _set_para_text(cell.paragraphs[0], text)

        def _cell_own_text(cell) -> str:
            """Direct paragraph text only — excludes text that lives in nested tables."""
            return ' '.join(p.text.strip() for p in cell.paragraphs if p.text.strip())

        def _collect_all_tables(tables_list):
            """Collect all tables recursively, including those nested inside cells."""
            result = []
            for tbl in tables_list:
                result.append(tbl)
                for row in tbl.rows:
                    for cell in row.cells:
                        nested = [_DocxTable(e, doc) for e in cell._tc.findall(qn('w:tbl'))]
                        if nested:
                            result.extend(_collect_all_tables(nested))
            return result

        # doc.tables only returns top-level tables, missing nested ones used
        # for side-by-side layouts (e.g. "Quote Details" / "Client Details").
        all_tables_in_doc = _collect_all_tables(doc.tables)

        # ── Pass 1: regex scan for [Placeholder] style fields ────────────────
        # Many templates use [Client Name], [Street Address] etc. Match these
        # directly so we don't rely on AI spatial reasoning for simple fields.
        _BRACKET_PATTERNS = {
            "customer_name": re.compile(
                r'^\[\s*(client\s*name|customer\s*name|name)\s*\]$', re.I
            ),
            "customer_address": re.compile(
                r'^\[\s*(client\s*address|customer\s*address|street\s*address|address\s*line\s*1?|address)\s*\]$', re.I
            ),
            "quote_ref": re.compile(
                r'^\[\s*(quote\s*(no\.?|ref|number|#)|invoice\s*(no\.?|ref|number|#)|ref\s*(no\.?|#)?)\s*\]$', re.I
            ),
            "quote_date": re.compile(
                r'^\[\s*(date|quote\s*date|invoice\s*date|dd[/\-\.]mm[/\-\.]yyyy)\s*\]$', re.I
            ),
        }
        _JINJA_FOR_FIELD = {
            "customer_name": "{{ customer_name }}",
            "customer_address": "{{ customer_address }}",
            "quote_ref": "{{ quote_ref }}",
            "quote_date": "{{ quote_date }}",
        }
        regex_matched = set()

        def _try_regex_inject_para(para):
            txt = para.text.strip()
            for field, pat in _BRACKET_PATTERNS.items():
                if field not in regex_matched and pat.match(txt):
                    _set_para_text(para, _JINJA_FOR_FIELD[field])
                    regex_matched.add(field)
                    logger.info(f"Regex-matched field '{field}' in paragraph: {txt!r}")
                    return

        def _try_regex_inject_cell(cell):
            for para in cell.paragraphs:
                txt = para.text.strip()
                for field, pat in _BRACKET_PATTERNS.items():
                    if field not in regex_matched and pat.match(txt):
                        _set_para_text(para, _JINJA_FOR_FIELD[field])
                        regex_matched.add(field)
                        logger.info(f"Regex-matched field '{field}' in cell: {txt!r}")
                        return

        for p in doc.paragraphs:
            _try_regex_inject_para(p)
        for table in all_tables_in_doc:
            for row in table.rows:
                for cell in row.cells:
                    _try_regex_inject_cell(cell)

        # ── Pass 2: Build compact structure and call Gemini for remaining fields
        non_empty_paras = [p for p in doc.paragraphs if p.text.strip()]
        structure = {
            "paragraphs": [{"index": i, "text": p.text.strip()} for i, p in enumerate(non_empty_paras[:60])],
            "tables": [],
        }
        for ti, table in enumerate(all_tables_in_doc):
            table_data = {"index": ti, "rows": []}
            for ri, row in enumerate(table.rows[:20]):
                # Use only direct paragraph text so outer layout cells appear
                # empty and Gemini focuses on the inner cells with real content.
                cells = [{"col": ci, "text": _cell_own_text(cell)} for ci, cell in enumerate(row.cells)]
                table_data["rows"].append(cells)
            structure["tables"].append(table_data)

        business_info = {k: brand_dna.get(k) for k in (
            "business_name", "business_address", "contact_details", "bank_info", "vat_tax_status"
        ) if brand_dna.get(k)}

        prompt = TEMPLATE_MAP_PROMPT.format(
            structure=json.dumps(structure, indent=2),
            business_info=json.dumps(business_info, indent=2),
        )

        try:
            response = _generate_with_retry(prompt)
            field_map = json.loads(response.text)
        except Exception as e:
            logger.error(f"Gemini template mapping failed: {e}")
            _capture(e)
            return None

        def _inject_location(loc, placeholder: str):
            if not loc or not isinstance(loc, dict):
                return
            try:
                if loc.get("type") == "paragraph":
                    idx = loc["paragraph_index"]
                    if idx < len(non_empty_paras):
                        _set_para_text(non_empty_paras[idx], placeholder)
                elif loc.get("type") == "table":
                    cell = all_tables_in_doc[loc["table_index"]].rows[loc["row_index"]].cells[loc["col_index"]]
                    _set_cell_text(cell, placeholder)
            except (IndexError, KeyError, TypeError) as e:
                logger.warning(f"Failed to inject placeholder '{placeholder}': {e}")

        # ── Inject simple field placeholders (skip fields matched by regex) ──
        if "customer_name" not in regex_matched:
            _inject_location(field_map.get("customer_name"), "{{ customer_name }}")
        if "customer_address" not in regex_matched:
            _inject_location(field_map.get("customer_address"), "{{ customer_address }}")
        if "quote_ref" not in regex_matched:
            _inject_location(field_map.get("quote_ref"), "{{ quote_ref }}")
        if "quote_date" not in regex_matched:
            _inject_location(field_map.get("quote_date"), "{{ quote_date }}")
        _inject_location(field_map.get("subtotal_value_location"), "{{ subtotal }}")
        _inject_location(field_map.get("tax_amount_value_location"), "{{ tax_amount }}")
        _inject_location(field_map.get("grand_total_value_location"), "{{ grand_total }}")

        # ── Inject line items table loop ─────────────────────────────────────
        li_ti = field_map.get("line_items_table_index")
        logger.info(f"Gemini field_map: {json.dumps({k: v for k, v in field_map.items() if k != 'line_items'}, default=str)}")

        # Fallback: if Gemini couldn't locate the line items table, detect it ourselves
        if li_ti is None:
            li_ti = _detect_line_items_table(all_tables_in_doc)
            if li_ti is not None:
                logger.info(f"Gemini returned null for line_items_table_index — fallback detected table {li_ti}")
            else:
                logger.warning("Could not detect line items table by any method — line items will be missing from template")

        if li_ti is not None:
            try:
                li_table = all_tables_in_doc[int(li_ti)]
                header_rows = set(field_map.get("line_item_header_rows") or [0])

                # Rows that hold totals must not be deleted when we strip extra data rows.
                # Gemini sometimes places subtotal/tax/total inside the same table.
                protected_rows = set()
                for loc_key in ("subtotal_value_location", "tax_amount_value_location", "grand_total_value_location"):
                    loc = field_map.get(loc_key)
                    if (loc and isinstance(loc, dict) and loc.get("type") == "table"
                            and int(loc.get("table_index", -1)) == int(li_ti)):
                        protected_rows.add(loc["row_index"])

                data_row_indices = [
                    i for i in range(len(li_table.rows))
                    if i not in header_rows and i not in protected_rows
                ]
                logger.info(f"Line items table {li_ti}: {len(li_table.rows)} rows, header_rows={header_rows}, protected_rows={protected_rows}, data_rows={data_row_indices}")

                if not data_row_indices:
                    # Template has only a header row — add a blank row to host the loop tag
                    li_table.add_row()
                    data_row_indices = [len(li_table.rows) - 1]

                if data_row_indices:
                    first_idx = data_row_indices[0]
                    template_row = li_table.rows[first_idx]
                    num_cols = len(template_row.cells)

                    # Map columns intelligently from the header row
                    header_row_idx = min(header_rows) if header_rows else 0
                    header_cells = li_table.rows[header_row_idx].cells if li_table.rows else []
                    col_field_map = _map_line_item_columns(header_cells)
                    # Positional fallback when header detection yields nothing
                    if not col_field_map:
                        _POSITIONAL = {0: 'description', 1: 'qty', 2: 'unit_price', 3: 'total'}
                        col_field_map = {ci: f for ci, f in _POSITIONAL.items() if ci < num_cols}
                        logger.info(f"Header detection empty — using positional column fallback")
                    logger.info(f"Line items column field map (from headers): {col_field_map}")

                    _FIELD_JINJA = {
                        'description': '{{ item.description }}',
                        'qty': '{{ item.qty }}',
                        'unit': '{{ item.unit_str }}',
                        'unit_price': '{{ item.unit_price_str }}',
                        'total': '{{ item.total_str }}',
                    }
                    # Set the template row to field expressions ONLY — no loop control tag.
                    # docxtpl replaces the ENTIRE <w:tr> that contains {%tr for %}, discarding
                    # all cell content. The for/endfor tags must live in their own dedicated
                    # rows so the field-expression row is preserved and repeated.
                    for ci in range(num_cols):
                        field = col_field_map.get(ci)
                        jinja_expr = _FIELD_JINJA.get(field, '') if field else ''
                        _set_cell_text(template_row.cells[ci], jinja_expr)

                    # Delete extra data rows (keep only the first as the loop template)
                    for ri in reversed(data_row_indices[1:]):
                        row_elem = li_table.rows[ri]._tr
                        li_table._tbl.remove(row_elem)

                    # Insert a {%tr for %} row immediately before the template row.
                    # add_row() appends at the end; addprevious() moves it into position.
                    for_row = li_table.add_row()
                    _set_cell_text(for_row.cells[0], "{%tr for item in line_items %}")
                    for ci in range(1, num_cols):
                        _set_cell_text(for_row.cells[ci], "")
                    template_row._tr.addprevious(for_row._tr)

                    # Insert a {%tr endfor %} row immediately after the template row.
                    endfor_row = li_table.add_row()
                    _set_cell_text(endfor_row.cells[0], "{%tr endfor %}")
                    for ci in range(1, len(endfor_row.cells)):
                        _set_cell_text(endfor_row.cells[ci], "")
                    template_row._tr.addnext(endfor_row._tr)

                    logger.info(f"Line items loop injected: {num_cols} columns, col_map={col_field_map}")

            except (IndexError, TypeError) as e:
                logger.warning(f"Failed to inject line items loop: {e}")

        # ── Final sweep: clear any remaining [Placeholder] bracket text ─────
        # Catches fields not in the regex/Gemini map: [Town, Postcode],
        # [client@email.com], [07xxx xxx xxx], etc.
        # [DD/MM/YYYY] on a second occurrence → valid_until placeholder.
        _PH_RE = re.compile(r'^\s*\[.+\]\s*$')
        _DATE_PH_RE = re.compile(r'^\s*\[DD[/\-.]MM[/\-.]YYYY\]\s*$', re.I)
        cleared = [0]

        def _sweep_para(para):
            txt = para.text.strip()
            if not txt:
                return
            if _DATE_PH_RE.match(txt):
                _set_para_text(para, "{{ valid_until }}")
                cleared[0] += 1
            elif _PH_RE.match(txt):
                _set_para_text(para, "")
                cleared[0] += 1

        for p in doc.paragraphs:
            _sweep_para(p)
        for table in all_tables_in_doc:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        _sweep_para(para)

        if cleared[0]:
            logger.info(f"Placeholder sweep cleared {cleared[0]} remaining bracket field(s)")

        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output.getvalue()

    @staticmethod
    def extract_brand_dna_from_xlsx(xlsx_path: str) -> dict | None:
        """
        Reads an XLSX template, builds a text dump of all cell values, and asks
        Gemini to extract Brand DNA. Returns the same dict shape as
        extract_brand_dna_from_blank (minus logo_base64, which XLSX cannot provide).
        """
        try:
            import openpyxl as _openpyxl

            try:
                wb = _openpyxl.load_workbook(xlsx_path, read_only=True, data_only=True)
            except Exception as e:
                logger.error(f"Failed to open XLSX for brand DNA extraction: {e}")
                return None

            lines = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                for row in ws.iter_rows():
                    for cell in row:
                        val = cell.value
                        if val is not None and str(val).strip():
                            lines.append(f"{sheet_name}!{cell.coordinate}: '{val}'")
            wb.close()

            cell_dump = "\n".join(lines)
            logger.info(f"XLSX cell dump: {len(lines)} non-empty cells, {len(cell_dump)} chars")

            if len(lines) < 3:
                logger.warning("XLSX text extraction yielded very little content")

            prompt = XLSX_DNA_PROMPT + f"\n\n--- Template cell values ---\n{cell_dump}"
            response = _generate_with_retry(prompt)

            try:
                raw = response.text or ""
            except Exception as e:
                logger.error(f"response.text raised: {e}", exc_info=True)
                return None

            if raw.startswith("```"):
                raw = re.sub(r"^```[a-zA-Z]*\n?", "", raw)
                raw = re.sub(r"\n?```$", "", raw)

            dna = json.loads(raw.strip())
            logger.info(f"XLSX brand DNA extracted: business_name={dna.get('business_name')!r}")
            return dna

        except RateLimitError:
            raise
        except Exception as e:
            logger.error(f"Failed to extract brand DNA from XLSX: {e}", exc_info=True)
            _capture(e)
            return None

    @staticmethod
    def build_xlsx_field_mapping(xlsx_path: str, brand_dna: dict) -> dict | None:
        """
        Reads an XLSX template and asks Gemini to identify the cell addresses for
        all variable quote fields. Returns a mapping dict or None on failure.
        """
        try:
            import openpyxl as _openpyxl

            try:
                wb = _openpyxl.load_workbook(xlsx_path, read_only=True, data_only=True)
            except Exception as e:
                logger.error(f"Failed to open XLSX for field mapping: {e}")
                return None

            lines = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                for row in ws.iter_rows():
                    for cell in row:
                        val = cell.value
                        if val is not None and str(val).strip():
                            lines.append(f"{sheet_name}!{cell.coordinate}: '{val}'")
            wb.close()

            cell_dump = "\n".join(lines)

            business_info = {k: brand_dna.get(k) for k in (
                "business_name", "business_address", "contact_details", "bank_info"
            ) if brand_dna.get(k)}

            prompt = XLSX_MAP_PROMPT.format(
                cell_dump=cell_dump,
                business_info=json.dumps(business_info, indent=2),
            )
            response = _generate_with_retry(prompt)

            try:
                raw = response.text or ""
            except Exception as e:
                logger.error(f"response.text raised during XLSX mapping: {e}", exc_info=True)
                return None

            if raw.startswith("```"):
                raw = re.sub(r"^```[a-zA-Z]*\n?", "", raw)
                raw = re.sub(r"\n?```$", "", raw)

            mapping = json.loads(raw.strip())
            logger.info(f"XLSX field mapping built: {mapping}")
            return mapping

        except RateLimitError:
            raise
        except Exception as e:
            logger.error(f"Failed to build XLSX field mapping: {e}", exc_info=True)
            _capture(e)
            return None
